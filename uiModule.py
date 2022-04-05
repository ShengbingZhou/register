# built-in package
import os
import shutil
import datetime

# pyside2 package
from PySide2.QtWidgets import QWidget, QAbstractItemView, QMessageBox, QMenu, QAction, QFileDialog, QProgressDialog
from PySide2.QtCore import Qt, Slot, QItemSelectionModel, QSize, QEvent, QDir, QFile, QUrl, QDir
from PySide2.QtGui import QStandardItemModel, QStandardItem, QIcon, QColor
from PySide2.QtSql import QSqlDatabase, QSqlTableModel, QSqlQueryModel, QSqlRecord, QSqlQuery
from PySide2.QtXmlPatterns import QXmlQuery, QXmlSerializer, QXmlResultItems
from PySide2.QtXml import QDomDocument, QDomNodeList

# lxml package
from lxml import etree    

# python-docx package
from docx import Document, oxml, shared
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

# local package
from ui.Module import Ui_ModuleWindow
from RegisterConst import RegisterConst
from QSqlQueryBfTableModel import QSqlQueryBfTableModel
from QSqlHighlightTableModel import QSqlHighlightTableModel, QRegValueDisplayDelegate
from QRegDebugTableModel import QRegDebugTableModel

class uiModuleWindow(QWidget):
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_ModuleWindow()
        self.ui.setupUi(self)
        self.ui.pbSetColumns.setVisible(False)
        self.ui.pbReadAll.setVisible(False)
        self.ui.pbReadSelected.setVisible(False)
        self.ui.pbWriteAll.setVisible(False)
        self.ui.treeView.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.ui.treeView.installEventFilter(self)
        self.ui.tableView.setAlternatingRowColors(True)
        self.ui.tableViewReg.setAlternatingRowColors(True)
        self.ui.tableViewReg.setVisible(False)
        with open (RegisterConst.StyleFile) as file:
            style = file.read()
        self.setStyleSheet(style)
        
        self.moduleIcon = QIcon('icon/module32.png')
        self.regMapIcon = QIcon('icon/regmap32.png')
        self.regIcon = QIcon('icon/reg32.png')
        self.bfIcon = QIcon('icon/bf32.png')
        self.bfenumIcon = QIcon('icon/bfenum32.png')
        
        self.ui.pbAddMemMap.setIcon(QIcon('icon/add32.png'))
        self.ui.pbAddRegMap.setIcon(QIcon('icon/add32.png'))
        self.ui.pbAddReg.setIcon(QIcon('icon/add32.png'))
        self.ui.pbAddBf.setIcon(QIcon('icon/add32.png'))
        self.ui.pbAddBfEnum.setIcon(QIcon('icon/add32.png'))

        self.ui.pbReadAll.setIcon(QIcon('icon/rdall32.png'))
        self.ui.pbReadSelected.setIcon(QIcon('icon/rdsel32.png'))
        self.ui.pbWriteAll.setIcon(QIcon('icon/wrall32.png'))
        
        # default as design view, no matter it is a new module or just opened module
        self.view = RegisterConst.DesignView
        
        # default as new module
        self.newModule = True
        self.fileName = ''
        self.newFileName = ''

        # default value
        self.__treeViewCurrentTable = None # selected table name by treeView
        self.__treeViewCurrentRow = None   # curent row index on treeView
        self.__regMapTypeIndex = None      # regmap table 'Type' column index which should be hidden, while this column should be visible for other tables
        self.__bfValueIndex = None         # bitfield table 'Value' column index which should be hidden, while this column should be visible for other tables
        return
    
    def eventFilter(self, obj, event):
        if obj == self.ui.treeView:
            if  event.type() == QEvent.KeyPress:
                if event.key() == Qt.Key_Delete:
                    self.do_delete_triggered()
        return super(uiModuleWindow, self).eventFilter(obj, event)
    
    def closeEvent(self, event):
        if self.newModule == True:
            reply = QMessageBox.information(self, "Save register ?", "Register design is created but never saved, do you want to save?", QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.mainWindow.on_actionSave_triggered()
        self.conn.close()
        if os.path.isfile(self.newFileName):
            os.remove(self.newFileName)
        event.accept()

    def setMainWindow(self, mainWindow):
        self.mainWindow = mainWindow

    def setView(self, view):
        if (view != self.view):
            self.view = view
            isDebugView = (self.view == RegisterConst.DebugView)
            self.ui.pbAddMemMap.setVisible(not isDebugView)
            self.ui.pbAddRegMap.setVisible(not isDebugView)
            self.ui.pbAddReg.setVisible(not isDebugView)
            self.ui.pbAddBf.setVisible(not isDebugView)
            self.ui.pbAddBfEnum.setVisible(not isDebugView)
            self.ui.labelDescription.setVisible(not isDebugView)
            self.ui.tableView.setAlternatingRowColors(not isDebugView)
            self.ui.tableView.setVisible(isDebugView or ((not isDebugView) and (self.__treeViewCurrentTable != "Register")))
            self.ui.tableViewReg.setVisible((not isDebugView) and (self.__treeViewCurrentTable == "Register"))
            self.ui.pbReadAll.setVisible(isDebugView)
            self.ui.pbReadSelected.setVisible(isDebugView)
            self.ui.pbWriteAll.setVisible(isDebugView)
            if (self.view == RegisterConst.DebugView):
                self.setupDebugViewModels()    
            self.do_treeView_currentChanged(self.ui.treeView.selectedIndexes().pop(), None)

    def getNewRowAndDisplayOrder(self, model, row, maxOrder):
        if row == -1 or row >= model.rowCount(): # append as last one in current model
            if model.rowCount() == 0: # no record in current model
                order = 0 if maxOrder == '' else int(maxOrder) + 1
                newRow = 0
            else:
                order = model.record(model.rowCount() - 1).value("DisplayOrder") + 1
                newRow = model.rowCount()
        else:
            order = model.record(row).value("DisplayOrder")
            newRow = row
        return newRow, order

    def newInfoRow(self, model, date):
        r = model.record()
        r.remove(r.indexOf('id'))
        r.setValue("Version", RegisterConst.Version)
        r.setValue("Author", os.getlogin())
        r.setValue("LastUpdateDate", date)
        model.insertRecord(-1, r)
        r = model.record(model.rowCount() - 1)
        return r

    def newMemMapRow(self, model, row):
        query = QSqlQuery(self.conn)
        query.exec_("SELECT max(DisplayOrder) FROM MemoryMap")
        query.next()
        maxOrder = query.record().value(0)
        newRow, order = self.getNewRowAndDisplayOrder(model, row, maxOrder)
        query.exec_("UPDATE MemoryMap SET DisplayOrder=DisplayOrder+1 WHERE DisplayOrder>=%s"%(order))

        query.exec_("INSERT INTO MemoryMap (DisplayOrder, OffsetAddress) VALUES ('%s', '%s')"%(order, 0x0000))

        query.exec_("SELECT max(id) FROM MemoryMap")
        query.next()
        id = query.record().value(0)
        query.exec_("UPDATE MemoryMap SET Name='MemoryMap%s' WHERE id=%s"%(id, id))
        model.select()
        r = model.record(newRow)
        return r

    def newRegMapRow(self, model, memMapId, row, type = RegisterConst.RegMap):
        query = QSqlQuery(self.conn)
        query.exec_("SELECT max(DisplayOrder) FROM RegisterMap")
        query.next()
        maxOrder = query.record().value(0)
        newRow, order = self.getNewRowAndDisplayOrder(model, row, maxOrder)
        query.exec_("UPDATE RegisterMap SET DisplayOrder=DisplayOrder+1 WHERE DisplayOrder>=%s"%(order))

        query.exec_("INSERT INTO RegisterMap (MemoryMapId, DisplayOrder, OffsetAddress, Description, Type) "\
                    "VALUES ('%s', '%s', '%s', '%s', '%s')"%(memMapId, order, 0x0000, "This is no name register map", type))

        query.exec_("SELECT max(id) FROM RegisterMap")
        query.next()
        id = query.record().value(0)
        if type == RegisterConst.RegMap:
            query.exec_("UPDATE RegisterMap SET Name='RegMap%s' WHERE id=%s"%(id, id))
        else:
            query.exec_("UPDATE RegisterMap SET Name='RegMod%s' WHERE id=%s"%(id, id))
        model.select()
        r = model.record(newRow)
        return r

    def newRegRow(self, model, regMapId, OffsetAddress, Width, row):
        query = QSqlQuery(self.conn)
        query.exec_("SELECT max(DisplayOrder) FROM Register")
        query.next()
        maxOrder = query.record().value(0)
        newRow, order = self.getNewRowAndDisplayOrder(model, row, maxOrder)
        query.exec_("UPDATE Register SET DisplayOrder=DisplayOrder+1 WHERE DisplayOrder>=%s"%(order))

        query.exec_("INSERT INTO Register (RegisterMapId, DisplayOrder, OffsetAddress, Description, Width) " \
                    "VALUES ('%s', '%s', '%s', '%s', '%s')"%(regMapId, order, OffsetAddress, "This is no name register", Width))

        query.exec_("SELECT max(id) FROM Register")
        query.next()
        id = query.record().value(0)
        query.exec_("UPDATE Register SET Name='Reg%s' WHERE id=%s"%(id, id))
        model.select()
        r = model.record(newRow)
        return r

    def newBfRow(self, model, regId, Width, row):  
        query = QSqlQuery(self.conn)
        query.exec_("SELECT max(DisplayOrder) FROM Bitfield")
        query.next()
        maxOrder = query.record().value(0)
        newRow, order = self.getNewRowAndDisplayOrder(model, row, maxOrder)
        query.exec_("UPDATE Bitfield SET DisplayOrder=DisplayOrder+1 WHERE DisplayOrder>=%s"%(order))

        query.exec_("INSERT INTO Bitfield (RegisterId, DisplayOrder, RegisterOffset, Description, Width, Access, DefaultValue, Value) " \
                    "VALUES ('%s', '%s', '%s', '%s', '%s', '%s', '%s', '%s')"%(regId, order, 0, "This is no name bitfield", Width, 'rw', 0, 0))

        query.exec_("SELECT max(id) FROM Bitfield")
        query.next()
        id = query.record().value(0)
        query.exec_("UPDATE Bitfield SET Name='Bf%s' WHERE id=%s"%(id, id))
        model.select()
        r = model.record(newRow)
        return r
    
    def newBfEnumRow(self, model, bfId, row):
        query = QSqlQuery(self.conn)
        query.exec_("SELECT max(DisplayOrder) FROM BitfieldEnum")
        query.next()
        maxOrder = query.record().value(0)
        newRow, order = self.getNewRowAndDisplayOrder(model, row, maxOrder)
        query.exec_("UPDATE BitfieldEnum SET DisplayOrder=DisplayOrder+1 WHERE DisplayOrder>=%s"%(order))

        query.exec_("INSERT INTO BitfieldEnum (BitfieldId, DisplayOrder, Description, Value) " \
                    "VALUES ('%s', '%s', '%s', '%s')"%(bfId, order, "This is no name enum", 0))

        query.exec_("SELECT max(id) FROM BitfieldEnum")
        query.next()
        id = query.record().value(0)
        query.exec_("UPDATE BitfieldEnum SET Name='BfEnum%s' WHERE id=%s"%(id, id))
        model.select()
        r = model.record(newRow)
        return r
    
    def newDatabase(self):
        # create temp database
        now = datetime.datetime.now().strftime('%Y_%m_%d_%H_%M_%S_%f')
        newName = "__%s%s"%(now, RegisterConst.DesignFileExt)
        newName = QDir.homePath() + "/.reg/" + newName        
        shutil.copyfile("template/module_template.db", newName)
        self.newFileName = newName
        self.newModule = True
        
        # open new database
        self.regDebugModels = [] # debug model is a list, each member is mapped to a regmap
        if self.setupDesignViewModels(newName):
            self.newInfoRow(self.infoTableModel, now).value("id")                 # create info   row0
            memMap0Id = self.newMemMapRow(self.memMapTableModel, -1).value("id")  # create memMap row0
            regMap0Id = self.newRegMapRow(self.regMapTableModel, memMap0Id, -1).value("id") # create regMap row0
            reg0Id = self.newRegRow(self.regTableModel, regMap0Id, 0, 32, -1).value("id")   # create register row0
            bf0Id = self.newBfRow(self.bfTableModel, reg0Id, 8, -1).value("id") # create bitField row0
            self.newBfEnumRow(self.bfEnumTableModel, bf0Id, -1).value("id")     # create bitFieldeEnum row0
            self.setupTreeView()
            self.regMapTableModel.dataChanged.connect(self.do_tableView_dataChanged)
            self.regTableModel.dataChanged.connect(self.do_tableView_dataChanged)
            self.bfTableModel.dataChanged.connect(self.do_tableView_dataChanged)
            self.bfEnumTableModel.dataChanged.connect(self.do_tableView_dataChanged)
        else:
            return False
        return True
        
    def openDatabase(self, fileName):
        # create temp database
        now = datetime.datetime.now().strftime('%Y_%m_%d_%H_%M_%S_%f')
        newName = "__%s%s"%(now, RegisterConst.DesignFileExt)
        newName = QDir.homePath() + "/.reg/" + newName  
        shutil.copyfile(fileName, newName)
        self.newFileName = newName
        self.fileName = fileName
        self.newModule = False

        # open existing database
        self.regDebugModels = [] # debug model is a list, each member is mapped to a regmap
        if self.setupDesignViewModels(newName):
            self.setupTreeView()
            self.regMapTableModel.dataChanged.connect(self.do_tableView_dataChanged)
            self.regTableModel.dataChanged.connect(self.do_tableView_dataChanged)
            self.bfTableModel.dataChanged.connect(self.do_tableView_dataChanged)
            self.bfEnumTableModel.dataChanged.connect(self.do_tableView_dataChanged)     
        else:
            return False
        return True
    
    def saveDatabase(self):
        fileName = ''
        if self.newModule == True:
            fileName, filterUsed = QFileDialog.getSaveFileName(self, "Save register file", QDir.homePath(), "Register File (*%s)"%RegisterConst.DesignFileExt)
            if fileName !='':
                if os.path.exists(fileName):                                        
                    os.remove(fileName)
                else:
                    f_name, f_ext = os.path.splitext(fileName)
                    # add .reg when saving new file
                    if f_ext != RegisterConst.DesignFileExt:
                        fileName += RegisterConst.DesignFileExt
                shutil.copy(self.newFileName, fileName)
                self.fileName = fileName
                self.newModule = False
        else:
            if self.newFileName != '':
                if os.path.exists(fileName):
                    os.remove(fileName)                  
                shutil.copy(self.newFileName, self.fileName)
                fileName = self.fileName
        return fileName    

    def importYodaSp1(self, fileName):
        # create temp database
        now = datetime.datetime.now().strftime('%Y_%m_%d_%H_%M_%S_%f')
        newName = "__%s%s"%(now, RegisterConst.DesignFileExt)
        newName = QDir.homePath() + "/.reg/" + newName   
        shutil.copyfile("template/module_template.db", newName)
        self.newFileName = newName
        self.newModule = True  

        # import .sp1 file
        self.regDebugModels = [] # debug model is a list, each member is mapped to a regmap
        if self.setupDesignViewModels(newName):
            infoId = self.newInfoRow(self.infoTableModel, now).value("id")

            # find root
            sp1 = etree.parse(fileName)
            root = sp1.getroot()
    
            # reset var
            memMapDisplayOrder = 0
            regMapDisplayOrder = 0
            regDisplayOrder = 0
            bfDisplayOrder = 0

            # find memory map
            memMapNodes = root.findall('MemoryMap')
            if len(memMapNodes) == 0:
                QMessageBox.warning(self, "Error", "Unable to find memory map", QMessageBox.Yes)
                if os.path.isfile(self.newFileName):
                    os.remove(self.newFileName)
                return False

            # start to import
            query = QSqlQuery(self.conn)
            for memMap in memMapNodes: # only one memory map in yoda
                # add memory record
                memMapAddr = root.find("Properties/Address").text.lower().replace("'h", "0x").replace("'d", "")
                query.exec_("INSERT INTO MemoryMap (DisplayOrder, OffsetAddress, Name) VALUES ('%s', '%s', '%s')"%(memMapDisplayOrder, memMapAddr, "MemoryMap"))
                query.exec_("SELECT max(id) FROM MemoryMap")
                query.next()
                memMapId = query.record().value(0)
                memMapDisplayOrder += 1

                # import regmap/reg/bf
                bfNodes = memMap.findall("BitFields/BitField")
                bfUIDs  = memMap.findall("BitFields/BitField/UID")
                regMapNodes = memMap.findall("RegisterMaps/RegisterMap")

                # prepare progress dialog
                dlgProgress = QProgressDialog("Importing %s ..."%fileName, "Cancel", 0, len(regMapNodes), self)
                dlgProgress.setWindowTitle("Importing...")
                dlgProgress.setWindowModality(Qt.WindowModal)

                for i in range(len(regMapNodes)):
                    regMapNode = regMapNodes[i]
                    regMapName = regMapNode.find("Name").text
                    regMapDesc = regMapNode.find("Description").text
                    regMapAddr = regMapNode.find("Address").text.lower().replace("'h", "0x").replace("'d", "")
                    query.exec_("INSERT INTO RegisterMap (MemoryMapId, DisplayOrder, Name, Description, OffsetAddress) " \
                                "VALUES ('%s', '%s', '%s', '%s', '%s')"%(memMapId, regMapDisplayOrder, regMapName, regMapDesc, regMapAddr))
                    query.exec_("SELECT max(id) FROM RegisterMap")
                    query.next()
                    regMapId = query.record().value(0)
                    regMapDisplayOrder += 1

                    dlgProgress.setLabelText("Importing register map '%s' from %s "%(regMapName, fileName))
                    dlgProgress.setValue(i)
                 
                    regNodes = regMapNode.findall("Registers/Register")
                    for j in range(len(regNodes)):
                        regNode = regNodes[j]
                        regName = regNode.find("Name").text
                        regDesc = regNode.find("Description").text
                        regAddr = regNode.find("Address").text.lower().replace("'h", "0x").replace("'d", "")
                        regWidth = regNode.find("Width").text.lower().replace("'h", "0x").replace("'d", "")
                        query.exec_("INSERT INTO Register (RegisterMapId, DisplayOrder, Name, Description, OffsetAddress, Width) " \
                                    "VALUES ('%s', '%s', '%s', '%s', '%s', '%s')"%(regMapId, regDisplayOrder, regName, regDesc, regAddr, regWidth))
                        query.exec_("SELECT max(id) FROM Register")
                        query.next()
                        regId = query.record().value(0)
                        regDisplayOrder += 1

                        bfRefNodes = regNode.findall("BitFieldRefs/BitFieldRef")
                        for k in range(len(bfRefNodes)):
                            bfRefNode = bfRefNodes[k]
                            bfUID = bfRefNode.find("BF-UID").text
                            regOffset  = bfRefNode.find("RegOffset").text.lower().replace("'h", "0x").replace("'d", "")
                            bitOffset  = bfRefNode.find("BitOffset").text.lower().replace("'h", "0x").replace("'d", "")
                            sliceWidth = bfRefNode.find("SliceWidth").text.lower().replace("'h", "0x").replace("'d", "")
                            for m in range(len(bfNodes)):
                                bfNode = bfNodes[m]
                                if bfUIDs[m].text == bfUID:
                                    bfName = bfNode.find("Name").text
                                    if bfName != "RESERVED":
                                        bfDesc = bfNode.find("Description").text
                                        bfDfValue = bfNode.find("DefaultValue").text.replace("'h", "0x").replace("'d", "").replace("'b", "").replace("'", "")
                                        # TODO: process default value to get sliced value
                                        query.exec_("INSERT INTO Bitfield (RegisterId, DisplayOrder, Name, Description, RegisterOffset, Width, DefaultValue) " \
                                                    "VALUES ('%s', '%s', '%s', '%s', '%s', '%s', '%s')"%(regId, bfDisplayOrder, bfName, bfDesc, regOffset, sliceWidth, bfDfValue))
                                        bfDisplayOrder += 1
                                    break
                dlgProgress.close()

            self.memMapTableModel.select()
            self.regMapTableModel.select()
            self.regTableModel.select()
            self.bfTableModel.select()
            self.setupTreeView()
            self.regMapTableModel.dataChanged.connect(self.do_tableView_dataChanged)
            self.regTableModel.dataChanged.connect(self.do_tableView_dataChanged)
            self.bfTableModel.dataChanged.connect(self.do_tableView_dataChanged)
            self.bfEnumTableModel.dataChanged.connect(self.do_tableView_dataChanged)

        return True    
    
    def importIpxact(self, fileName):
        # create temp database
        now = datetime.datetime.now().strftime('%Y_%m_%d_%H_%M_%S_%f')
        newName = "__%s%s"%(now, RegisterConst.DesignFileExt)
        newName = QDir.homePath() + "/.reg/" + newName   
        shutil.copyfile("template/module_template.db", newName)
        self.newFileName = newName
        self.newModule = True  

        # import ipxact xml file
        self.regDebugModels = [] # debug model is a list, each member is mapped to a regmap
        if self.setupDesignViewModels(newName):
            infoId = self.newInfoRow(self.infoTableModel, now).value("id")

            # find root
            sp1 = etree.parse(fileName)
            root = sp1.getroot()
            
            # reset var
            memMapDisplayOrder = 0
            regMapDisplayOrder = 0
            regDisplayOrder = 0
            bfDisplayOrder = 0            

            # find memory map
            memMapNodes = root.findall('ipxact:memoryMaps/ipxact:memoryMap', root.nsmap)
            if len(memMapNodes) == 0:
                QMessageBox.warning(self, "Error", "Unable to find memory map", QMessageBox.Yes)
                if os.path.isfile(self.newFileName):
                    os.remove(self.newFileName)
                return False

            # start to import
            query = QSqlQuery(self.conn)
            for memMap in range(len(memMapNodes)):
                memMapNode = memMapNodes[memMap]
                memMapName = memMapNode.find("ipxact:name", root.nsmap).text

                # add memory record
                query.exec_("INSERT INTO MemoryMap (DisplayOrder, OffsetAddress, Name) VALUES ('%s', '%s', '%s')"%(memMapDisplayOrder, 0x0000, memMapName))
                query.exec_("SELECT max(id) FROM MemoryMap")
                query.next()
                memMapId = query.record().value(0)
                memMapDisplayOrder += 1

                # get regmap nodes
                regMapNodes = memMapNode.findall("ipxact:addressBlock", root.nsmap)

                # prepare progress dialog
                dlgProgress = QProgressDialog("Importing %s ..."%fileName, "Cancel", 0, len(regMapNodes), self)
                dlgProgress.setWindowTitle("Importing...")
                dlgProgress.setWindowModality(Qt.WindowModal)

                for i in range(len(regMapNodes)):
                    regMapNode = regMapNodes[i]
                    regMapName = regMapNode.find("ipxact:name", root.nsmap).text

                    n = regMapNode.find("ipxact:description", root.nsmap)
                    regMapDesc = "" if n == None else n.text

                    regMapAddr = regMapNode.find("ipxact:baseAddress", root.nsmap)
                    params = regMapNode.findall("ipxact:parameters/ipxact:parameter", root.nsmap)
                    for param in params:
                        if param.find("ipxact:name", root.nsmap).text == regMapAddr.text:
                            regMapAddr = param.find("ipxact:value", root.nsmap)
                            break
                    regMapAddr = regMapAddr.text.lower().replace("'h", "0x").replace("'d", "")
                    query.exec_("INSERT INTO RegisterMap (MemoryMapId, DisplayOrder, Name, Description, OffsetAddress) " \
                                "VALUES ('%s', '%s', '%s', '%s', '%s')"%(memMapId, regMapDisplayOrder, regMapName, regMapDesc, regMapAddr))
                    query.exec_("SELECT max(id) FROM RegisterMap")
                    query.next()
                    regMapId = query.record().value(0)
                    regMapDisplayOrder += 1

                    dlgProgress.setLabelText("Importing register map '%s' from %s "%(regMapName, fileName))
                    dlgProgress.setValue(i)
                 
                    regNodes = regMapNode.findall("ipxact:register", root.nsmap)
                    for j in range(len(regNodes)):
                        regNode = regNodes[j]
                        regName = regNode.find("ipxact:name", root.nsmap).text

                        n = regNode.find("ipxact:description", root.nsmap)
                        regDesc = "" if n == None else n.text

                        regAddr = regNode.find("ipxact:addressOffset", root.nsmap).text.lower().replace("'h", "0x").replace("'d", "")
                        regWidth = regNode.find("ipxact:size", root.nsmap).text.lower().replace("'h", "0x").replace("'d", "")
                        query.exec_("INSERT INTO Register (RegisterMapId, DisplayOrder, Name, Description, OffsetAddress, Width) " \
                                    "VALUES ('%s', '%s', '%s', '%s', '%s', '%s')"%(regMapId, regDisplayOrder, regName, regDesc, regAddr, regWidth))
                        query.exec_("SELECT max(id) FROM Register")
                        query.next()
                        regId = query.record().value(0)
                        regDisplayOrder += 1

                        bfNodes = regNode.findall("ipxact:field", root.nsmap)
                        for k in range(len(bfNodes)):
                            bfNode = bfNodes[k]
                            bfName = bfNode.find("ipxact:name", root.nsmap).text

                            n = bfNode.find("ipxact:description", root.nsmap)
                            bfDesc = "" if n == None else n.text

                            regOffset = bfNode.find("ipxact:bitOffset", root.nsmap).text.lower().replace("'h", "0x").replace("'d", "")
                            bfWidth   = bfNode.find("ipxact:bitWidth", root.nsmap).text.lower().replace("'h", "0x").replace("'d", "")
                            
                            n = bfNode.find("ipxact:resets/ipxact:reset/ipxact:value", root.nsmap)
                            bfDefaultVal = 0 if n is None else n.text.lower().replace("'h", "0x").replace("'d", "")
                            
                            query.exec_("INSERT INTO Bitfield (RegisterId, DisplayOrder, Name, Description, RegisterOffset, Width, DefaultValue) " \
                                        "VALUES ('%s', '%s', '%s', '%s', '%s', '%s', '%s')"%(regId, bfDisplayOrder, bfName, bfDesc, regOffset, bfWidth, bfDefaultVal))
                            bfDisplayOrder += 1
                dlgProgress.close()

            self.memMapTableModel.select()
            self.regMapTableModel.select()
            self.regTableModel.select()
            self.bfTableModel.select()
            self.setupTreeView()
            self.regMapTableModel.dataChanged.connect(self.do_tableView_dataChanged)
            self.regTableModel.dataChanged.connect(self.do_tableView_dataChanged)
            self.bfTableModel.dataChanged.connect(self.do_tableView_dataChanged)
            self.bfEnumTableModel.dataChanged.connect(self.do_tableView_dataChanged)
        return True

    def exporIpxact(self):
        fileName, filterUsed = QFileDialog.getSaveFileName(self, "Export ipxact file", QDir.homePath(), "ipxact File (*.xml)")
        if fileName !='':
            f_name, f_ext = os.path.splitext(os.path.basename(fileName))
            # add .xml when saving ipxact file
            if f_ext != ".xml":
                fileName += ".xml"               
            ipxactFile = open(fileName, "w")
            ipxactFile.write("<ipxact:component xmlns:xsi=\"http://www.w3.org/2001/XMLSchema-instance\" xmlns:ipxact=\"http://www.accellera.org/XMLSchema/IPXACT/1685-2014\" xsi:schemaLocation=\"http://www.accellera.org/XMLSchema/IPXACT/1685-2014 http://www.accellera.org/XMLSchema/IPXACT/1685-2014/index.xsd\">\n")
            ipxactFile.write("  <ipxact:vendor>Register@ShengbingZhou</ipxact:vendor>\n")
            ipxactFile.write("  <ipxact:name>%s</ipxact:name>\n"%f_name)
            ipxactFile.write("  <ipxact:version>1.0</ipxact:version>\n")
            ipxactFile.write("  <ipxact:memoryMaps>\n")
            
            # memory map
            memoryMapQueryModel = QSqlQueryModel()
            memoryMapQueryModel.setQuery("SELECT * FROM MemoryMap", self.conn)
            for i in range(memoryMapQueryModel.rowCount()):
                memMapRecord = memoryMapQueryModel.record(i)
                ipxactFile.write("    <ipxact:memoryMap>\n")
                ipxactFile.write("      <ipxact:name>%s</ipxact:name>\n"%memoryMapQueryModel.value("Name"))
                
                # register map
                regMapQueryModel = QSqlQueryModel()
                regMapQueryModel.setQuery("SELECT * FROM RegisterMap WHERE memoryMapId=%s ORDER BY DisplayOrder ASC"%memMapRecord.value("id"), self.conn)
                for j in range(regMapQueryModel.rowCount()):
                    regMapRecord = regMapQueryModel.record(j)
                    ipxactFile.write("      <ipxact:addressBlock>\n")
                    ipxactFile.write("        <ipxact:name>%s</ipxact:name>\n"%(regMapRecord.value("Name")))
                    ipxactFile.write("        <ipxact:description>%s</ipxact:description>\n"%(regMapRecord.value("Description")))
                    ipxactFile.write("        <ipxact:baseAddress>%s</ipxact:baseAddress>\n"%(str(regMapRecord.value("OffsetAddress")).replace("0x", "'h")))
                    #ipxactFile.write("        <ipxact:width>%s</ipxact:width>\n"%(regMapRecord.value("Width")))
                    
                    # register
                    regQueryModel = QSqlQueryModel()
                    regQueryModel.setQuery("SELECT * FROM Register WHERE RegisterMapId=%s ORDER BY DisplayOrder ASC"%regMapRecord.value("id"), self.conn)
                    for k in range(regQueryModel.rowCount()):
                        regRecord = regQueryModel.record(k)
                        ipxactFile.write("        <ipxact:register>\n")
                        ipxactFile.write("          <ipxact:name>%s</ipxact:name>\n"%(regRecord.value("Name")))
                        ipxactFile.write("          <ipxact:description>%s</ipxact:description>\n"%(regRecord.value("Description")))
                        ipxactFile.write("          <ipxact:addressOffset>%s</ipxact:addressOffset>\n"%(str(regRecord.value("OffsetAddress")).replace("0x", "'h")))
                        ipxactFile.write("          <ipxact:size>%s</ipxact:size>\n"%(regRecord.value("Width")))

                        # bitfield
                        bfQueryModel = QSqlQueryModel()
                        bfQueryModel.setQuery("SELECT * FROM Bitfield WHERE RegisterId=%s ORDER BY DisplayOrder ASC"%regRecord.value("id"), self.conn)
                        for m in range(bfQueryModel.rowCount()):
                            bfRecord = bfQueryModel.record(m)
                            ipxactFile.write("          <ipxact:field>\n")
                            ipxactFile.write("            <ipxact:name>%s</ipxact:name>\n"%(bfRecord.value("Name")))
                            ipxactFile.write("            <ipxact:bitOffset>%s</ipxact:bitOffset>\n"%(bfRecord.value("RegisterOffset")))
                            ipxactFile.write("            <ipxact:bitWidth>%s</ipxact:bitWidth>\n"%(bfRecord.value("Width")))
                            ipxactFile.write("            <ipxact:access>%s</ipxact:access>\n"%("read-write"))
                            ipxactFile.write("            <ipxact:volatile>%s</ipxact:volatile>\n"%("true"))
                            ipxactFile.write("            <ipxact:resets>\n")
                            ipxactFile.write("              <ipxact:reset>\n")
                            ipxactFile.write("                <ipxact:value>%s</ipxact:value>\n"%(str(bfRecord.value("DefaultValue")).replace("0x", "'h")))
                            ipxactFile.write("              </ipxact:reset>\n")
                            ipxactFile.write("              </ipxact:reset>\n")
                            ipxactFile.write("            </ipxact:resets>\n")
                            ipxactFile.write("          </ipxact:field>\n")
                        ipxactFile.write("        </ipxact:register>\n")
                    ipxactFile.write("      </ipxact:addressBlock>\n")
                ipxactFile.write("    </ipxact:memoryMap>\n")
            ipxactFile.write("  </ipxact:memoryMaps>\n")
            ipxactFile.write("</ipxact:component>\n")
            ipxactFile.close()  
        return

    def exporDocx(self):
        fileName, filterUsed = QFileDialog.getSaveFileName(self, "Export Word file", QDir.homePath(), "Word File (*.docx)")
        if fileName !='':
            f_name, f_ext = os.path.splitext(os.path.basename(fileName))
            # add .docx
            if f_ext != ".docx":
                fileName += ".docx"
            docx = Document()
            docx.styles['Heading 1'].font.size = shared.Pt(11)
            docx.styles['Heading 2'].font.size = shared.Pt(10)
            docx.styles['Heading 3'].font.size = shared.Pt(9)
            docx.styles['Heading 4'].font.size = shared.Pt(8)
            docx.styles['Normal'].font.size = shared.Pt(8)
            
            # memory map
            memoryMapQueryModel = QSqlQueryModel()
            memoryMapQueryModel.setQuery("SELECT * FROM MemoryMap", self.conn)

            title = docx.add_heading('MemoryMap Table\n', level = 1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            fields = ['Name', 'Description']
            table = docx.add_table(rows=memoryMapQueryModel.rowCount() + 1, cols=len(fields), style='Table Grid')

            for i, row in enumerate(table.rows):
                for j, (cell, field) in enumerate(zip(row.cells, fields)):
                    if i == 0: # table header
                        cell.text = fields[j]
                        cell._tc.get_or_add_tcPr().append(oxml.parse_xml(r'<w:shd {} w:fill="c0c0c0"/>'.format(oxml.ns.nsdecls('w'))))
                    else:
                        memMapRecord = memoryMapQueryModel.record(i - 1)
                        if field == 'Name':
                            cell.text = memMapRecord.value("Name")
            docx.add_page_break()

            for i in range(memoryMapQueryModel.rowCount()):
                memMapRecord = memoryMapQueryModel.record(i)
                docx.add_heading('%s'%(memMapRecord.value("Name").upper()), level = 2)

                # register map
                regMapQueryModel = QSqlQueryModel()
                regMapQueryModel.setQuery("SELECT * FROM RegisterMap WHERE memoryMapId=%s ORDER BY DisplayOrder ASC"%memMapRecord.value("id"), self.conn)
                for j in range(regMapQueryModel.rowCount()):
                    regMapRecord = regMapQueryModel.record(j)
                    docx.add_heading('%s'%(regMapRecord.value("Name")), level = 3)
                    docx.add_paragraph('Description : %s'%(regMapRecord.value("Description")))
                    docx.add_paragraph('BaseAddress : %s'%(regMapRecord.value("OffsetAddress")))
                    #docx.add_paragraph('Width : %bits'%(regMapRecord.value("Width")))
            
                    # register
                    regQueryModel = QSqlQueryModel()
                    regQueryModel.setQuery("SELECT * FROM Register WHERE RegisterMapId=%s ORDER BY DisplayOrder ASC"%regMapRecord.value("id"), self.conn)

                    fields = ['Name', 'Address', 'Description']
                    table = docx.add_table(rows=regQueryModel.rowCount() + 1, cols=len(fields), style='Table Grid')
                    for r, row in enumerate(table.rows):
                        for c, (cell, field) in enumerate(zip(row.cells, fields)):
                            if r == 0:
                                cell.text = fields[c]
                                cell._tc.get_or_add_tcPr().append(oxml.parse_xml(r'<w:shd {} w:fill="c0c0c0"/>'.format(oxml.ns.nsdecls('w'))))
                            else:
                                regRecord = regQueryModel.record(r - 1)
                                if field == 'Name':
                                    cell.text = regRecord.value("Name")
                                if field == 'Address':
                                    cell.text = "%s"%regRecord.value("OffsetAddress")
                                if field == 'Description':
                                    cell.text = regRecord.value("Description")

                    for k in range(regQueryModel.rowCount()):
                        regRecord = regQueryModel.record(k)
                        docx.add_heading('%s'%(regRecord.value("Name")), level = 4)
                        docx.add_paragraph('Description : %s'%(regRecord.value("Description")))
                        docx.add_paragraph('Address : %s'%(regRecord.value("OffsetAddress")))

                        # bitfield
                        bfQueryModel = QSqlQueryModel()
                        bfQueryModel.setQuery("SELECT * FROM Bitfield WHERE RegisterId=%s ORDER BY DisplayOrder ASC"%regRecord.value("id"), self.conn)

                        fields = ['Name', 'Offset', 'Width', 'ResetValue', 'Description']
                        table = docx.add_table(rows=bfQueryModel.rowCount() + 1, cols=len(fields), style='Table Grid')
                        for r, row in enumerate(table.rows):
                            for c, (cell, field) in enumerate(zip(row.cells, fields)):
                                if r == 0:
                                    cell.text = fields[c]
                                    cell._tc.get_or_add_tcPr().append(oxml.parse_xml(r'<w:shd {} w:fill="c0c0c0"/>'.format(oxml.ns.nsdecls('w'))))
                                else:
                                    bfRecord = bfQueryModel.record(r - 1)
                                    if field == 'Name':
                                        cell.text = bfRecord.value("Name")
                                    if field == 'Offset':
                                        cell.text = "%s"%(bfRecord.value("RegisterOffset"))
                                    if field == 'Width':
                                        cell.text = "%s"%(bfRecord.value("Width"))
                                    if field == 'ResetValue':
                                        cell.text = "%s"%(bfRecord.value("DefaultValue"))
                                    if field == 'Description':
                                        cell.text = bfRecord.value("Description")
                    docx.add_page_break()
            docx.add_page_break()
            docx.save(fileName)
            QMessageBox.information(self, "Exporting docx", "Done!", QMessageBox.Yes)
        return

    def setupDesignViewModels(self, fileName):  
        self.conn = QSqlDatabase.addDatabase("QSQLITE", fileName)
        self.conn.setDatabaseName(fileName)
        if self.conn.open():
            # enable FK
            query = QSqlQuery(self.conn)
            query.prepare("PRAGMA foreign_keys = ON")
            query.exec_()

            # setup table models
            self.infoTableModel = QSqlTableModel(self, self.conn)
            self.infoTableModel.setEditStrategy(QSqlTableModel.OnFieldChange) 
            self.infoTableModel.setTable("info")
            self.infoTableModel.select()
            
            self.memMapTableModel = QSqlTableModel(self, self.conn)
            self.memMapTableModel.setEditStrategy(QSqlTableModel.OnFieldChange)  
            self.memMapTableModel.setTable("MemoryMap")
            self.memMapTableModel.setSort(self.memMapTableModel.fieldIndex("DisplayOrder"), Qt.AscendingOrder)
            self.memMapTableModel.select()

            self.regMapTableModel = QSqlHighlightTableModel(self, self.conn)
            self.regMapTableModel.setEditStrategy(QSqlTableModel.OnFieldChange)  
            self.regMapTableModel.setTable("RegisterMap")
            self.regMapTableModel.setSort(self.regMapTableModel.fieldIndex("DisplayOrder"), Qt.AscendingOrder)
            self.regMapTableModel.select()
            
            self.regTableModel = QSqlHighlightTableModel(self, self.conn)
            self.regTableModel.setEditStrategy(QSqlTableModel.OnFieldChange)  
            self.regTableModel.setTable("Register")
            self.regTableModel.setSort(self.regTableModel.fieldIndex("DisplayOrder"), Qt.AscendingOrder)
            self.regTableModel.select()
                        
            self.bfTableModel = QSqlHighlightTableModel(self, self.conn)
            self.bfTableModel.setEditStrategy(QSqlTableModel.OnFieldChange) 
            self.bfTableModel.setTable("Bitfield")
            self.bfTableModel.setSort(self.bfTableModel.fieldIndex("DisplayOrder"), Qt.AscendingOrder)
            self.bfTableModel.select()
            
            self.bfEnumTableModel = QSqlHighlightTableModel(self, self.conn)
            self.bfEnumTableModel.setEditStrategy(QSqlTableModel.OnFieldChange) 
            self.bfEnumTableModel.setTable("BitfieldEnum")
            self.bfEnumTableModel.setSort(self.bfEnumTableModel.fieldIndex("DisplayOrder"), Qt.AscendingOrder)
            self.bfEnumTableModel.select()
        else:
            QMessageBox.warning(self, "Error", "Failed to open %s"%fileName)  
            return False
        return True

    def setupTreeView(self):
        # create standard model for treeview          
        self.treeViewTableModel = QStandardItemModel()
        root = self.treeViewTableModel.invisibleRootItem()

        # memory map
        memoryMapQueryModel = QSqlQueryModel()
        memoryMapQueryModel.setQuery("SELECT id, Name FROM MemoryMap", self.conn)
        for i in range(memoryMapQueryModel.rowCount()):
            memMapRecord = memoryMapQueryModel.record(i)
            self.memoryMapItem = QStandardItem(self.moduleIcon, memMapRecord.value("name"))
            self.memoryMapItem.setData("MemoryMap", RegisterConst.NameRole)
            self.memoryMapItem.setData(memMapRecord.value("id"), RegisterConst.MemMapIdRole)
            root.appendRow(self.memoryMapItem)            

            # register map
            regMapQueryModel = QSqlQueryModel()
            regMapQueryModel.setQuery("SELECT id, Name, Type FROM RegisterMap WHERE memoryMapId=%s ORDER BY DisplayOrder ASC"%memMapRecord.value("id"), self.conn)
            for j in range(regMapQueryModel.rowCount()):
                regMapRecord = regMapQueryModel.record(j)
                regMapitem = QStandardItem(self.regMapIcon, regMapRecord.value("name"))
                regMapitem.setData("RegisterMap", RegisterConst.NameRole)
                regMapitem.setData(memMapRecord.value("id"), RegisterConst.MemMapIdRole)
                regMapitem.setData(regMapRecord.value("id"), RegisterConst.RegMapIdRole)
                regMapitem.setData(regMapRecord.value("Type"), RegisterConst.RegMapTypeRole)
                self.memoryMapItem.appendRow(regMapitem)
                if RegisterConst.recordExist(regMapRecord) == False:
                    regMapitem.setData(QColor('grey'), Qt.BackgroundColorRole)
            
                # register
                regQueryModel = QSqlQueryModel()
                regQueryModel.setQuery("SELECT id, Name FROM Register WHERE RegisterMapId=%s ORDER BY DisplayOrder ASC"%regMapRecord.value("id"), self.conn)
                for k in range(regQueryModel.rowCount()):
                    regRecord = regQueryModel.record(k)
                    regItem = QStandardItem(self.regIcon, regRecord.value("name"))
                    regItem.setData("Register", RegisterConst.NameRole)
                    regItem.setData(memMapRecord.value("id"), RegisterConst.MemMapIdRole)
                    regItem.setData(regMapRecord.value("id"), RegisterConst.RegMapIdRole)
                    regItem.setData(regRecord.value("id"), RegisterConst.RegIdRole)
                    regMapitem.appendRow(regItem)
                    if RegisterConst.recordExist(regRecord) == False:
                        regItem.setData(QColor('grey'), Qt.BackgroundColorRole)
                    
                    # bitfield
                    bfQueryModel = QSqlQueryModel()
                    bfQueryModel.setQuery("SELECT id, Name FROM Bitfield WHERE RegisterId=%s ORDER BY DisplayOrder ASC"%regRecord.value("id"), self.conn)
                    for m in range(bfQueryModel.rowCount()):
                        bfRecord = bfQueryModel.record(m)
                        bfItem = QStandardItem(self.bfIcon, bfRecord.value("name"))
                        bfItem.setData("Bitfield", RegisterConst.NameRole)
                        bfItem.setData(memMapRecord.value("id"), RegisterConst.MemMapIdRole)
                        bfItem.setData(regMapRecord.value("id"), RegisterConst.RegMapIdRole)
                        bfItem.setData(regRecord.value("id"), RegisterConst.RegIdRole)
                        bfItem.setData(bfRecord.value("id"), RegisterConst.BfIdRole)
                        regItem.appendRow(bfItem)
                        if RegisterConst.recordExist(bfRecord) == False:
                            bfItem.setData(QColor('grey'), Qt.BackgroundColorRole)                       

                        # bitfield enum
                        bfEnumQueryModel = QSqlQueryModel()
                        bfEnumQueryModel.setQuery("SELECT id, Name FROM BitfieldEnum WHERE BitfieldId=%s ORDER BY DisplayOrder ASC"%bfRecord.value("id"), self.conn)
                        for n in range(bfEnumQueryModel.rowCount()):
                            bfEnumRecord = bfEnumQueryModel.record(n)
                            bfEnumItem = QStandardItem(self.bfenumIcon, bfEnumRecord.value("name"))
                            bfEnumItem.setData("BitfieldEnum", RegisterConst.NameRole)
                            bfEnumItem.setData(memMapRecord.value("id"), RegisterConst.MemMapIdRole)
                            bfEnumItem.setData(regMapRecord.value("id"), RegisterConst.RegMapIdRole)
                            bfEnumItem.setData(regRecord.value("id"), RegisterConst.RegIdRole)
                            bfEnumItem.setData(bfRecord.value("id"), RegisterConst.BfIdRole)
                            bfEnumItem.setData(bfEnumRecord.value("id"), RegisterConst.BfEnumIdRole)
                            bfItem.appendRow(bfEnumItem)
                            if RegisterConst.recordExist(bfEnumRecord) == False:
                                bfEnumItem.setData(QColor('grey'), Qt.BackgroundColorRole)
        
        # show tree view nodes
        self.ui.treeView.setModel(self.treeViewTableModel)
        self.ui.treeView.expandAll()

        # connect slots
        treeViewSelectionModel = self.ui.treeView.selectionModel()
        treeViewSelectionModel.currentChanged.connect(self.do_treeView_currentChanged)

        self.ui.treeView.setContextMenuPolicy(Qt.CustomContextMenu)
        self.ui.treeView.customContextMenuRequested.connect(self.do_treeView_contextMenuRequested)
        
        # select memory map node
        memoryMapItemIndex = self.treeViewTableModel.indexFromItem(self.memoryMapItem)
        treeViewSelectionModel.select(memoryMapItemIndex, QItemSelectionModel.ClearAndSelect)
        self.do_treeView_currentChanged(memoryMapItemIndex, None)
    
    def setupDebugViewModels(self):
        # clear existing item
        for model in self.regDebugModels:
            model.clear()
        self.regDebugModels.clear()

        # register map
        regMapQueryModel = QSqlQueryModel()
        regMapQueryModel.setQuery("SELECT * FROM RegisterMap ORDER BY DisplayOrder ASC", self.conn)
        for i in range(regMapQueryModel.rowCount()):
            regMapRecord = regMapQueryModel.record(i)
            if RegisterConst.recordExist(regMapRecord) == True:
                debugModel = QRegDebugTableModel()
                debugModel.setRegMapId(regMapRecord.value("id"))
                debugModel.setHorizontalHeaderLabels(["Name", "Address", "Description", "Value"])
                self.regDebugModels.append(debugModel)
            
            # register
            regQueryModel = QSqlQueryModel()
            regQueryModel.setQuery("SELECT * FROM Register WHERE RegisterMapId=%s ORDER BY DisplayOrder ASC"%regMapRecord.value("id"), self.conn)
            for j in range(regQueryModel.rowCount()):
                regRecord = regQueryModel.record(j)
                if RegisterConst.recordExist(regRecord) == True:
                    regItem0 = QStandardItem(self.regIcon, regRecord.value("name"))
                    regItem1 = QStandardItem(str(regRecord.value("OffsetAddress")))
                    regItem2 = QStandardItem(regRecord.value("Description"))
                    regItem3 = QStandardItem(str(regRecord.value("Value")))
                    regItems = [regItem0, regItem1, regItem2, regItem3]
                    for r in regItems:
                        r.setData("Register", RegisterConst.NameRole)
                    debugModel.appendRow(regItems)

                # bitfield
                bfQueryModel = QSqlQueryModel()
                bfQueryModel.setQuery("SELECT Name, Value FROM Bitfield WHERE RegisterId=%s ORDER BY DisplayOrder ASC"%regRecord.value("id"), self.conn)
                for k in range(bfQueryModel.rowCount()):
                    bfRecord = bfQueryModel.record(k)
                    if RegisterConst.recordExist(bfRecord) == True:
                        bfItem0 = QStandardItem(" ")
                        bfItem1 = QStandardItem(" ")
                        bfItem2 = QStandardItem(bfRecord.value("name"))
                        bfItem3 = QStandardItem(str(bfRecord.value("Value")))
                        bfItems = [bfItem0, bfItem1, bfItem2, bfItem3]
                        for r in bfItems:
                            r.setData("Bitfield", RegisterConst.NameRole)                        
                        debugModel.appendRow(bfItems)

    @Slot('QItemSelection', 'QItemSelection')
    def do_tableView_selectionChanged(self, selected, deselected):
        if self.view == RegisterConst.DesignView:
            if self.__treeViewCurrentTable == "Register":                            
                tableViewCurrents = self.ui.tableViewReg.selectionModel().selectedIndexes()                
            else:
                tableViewCurrents = self.ui.tableView.selectionModel().selectedIndexes()
            tableViewCurrent = None if len(tableViewCurrents) == 0 else tableViewCurrents[0]
            if tableViewCurrent != None:
                treeViewCurrent = self.ui.treeView.selectedIndexes()[0]
                if tableViewCurrent.row() != self.__treeViewCurrentRow:
                    sibling = treeViewCurrent.sibling(tableViewCurrent.row(), 0)
                    self.ui.treeView.selectionModel().setCurrentIndex(sibling, QItemSelectionModel.ClearAndSelect)                
        return
    
    @Slot()
    def do_treeView_contextMenuRequested(self, point):
        # do not allow to edit design by context menu in debug view
        if self.view == RegisterConst.DebugView:
            return
        
        # allow to edit by context menu in design view
        index = self.ui.treeView.indexAt(point)
        tableName = str(index.data(RegisterConst.NameRole))
        
        self.treeViewPopMenu = QMenu(self)
        addMemMapAction = QAction("+ MemMap", self)
        addRegMapAction = QAction("+ RegMap", self)
        addRegModAction = QAction("+ RegMod", self)
        addRegAction = QAction("+ Reg", self)
        addBfAction = QAction("+ Bitfield", self)
        addBfEnumAction = QAction("+ Bitfield Enum", self)
        
        if tableName == "MemoryMap":
            self.treeViewPopMenu.addAction(addMemMapAction)
            addMemMapAction.triggered.connect(self.on_pbAddMemMap_clicked)
            self.treeViewPopMenu.addAction(addRegMapAction)
            addRegMapAction.triggered.connect(self.on_pbAddRegMap_clicked)
            self.treeViewPopMenu.addAction(addRegModAction)
            addRegModAction.triggered.connect(self.do_pbAddRegMod_clicked)
        elif tableName == "RegisterMap":
            self.treeViewPopMenu.addAction(addRegMapAction)
            addRegMapAction.triggered.connect(self.on_pbAddRegMap_clicked)
            self.treeViewPopMenu.addAction(addRegModAction)
            addRegModAction.triggered.connect(self.do_pbAddRegMod_clicked)
            regMapType = index.data(RegisterConst.RegMapTypeRole)
            regMapType = RegisterConst.RegMap if regMapType == None or regMapType == '' else int(regMapType) # default as regmap if not set
            if regMapType == RegisterConst.RegMap:
                self.treeViewPopMenu.addAction(addRegAction)
                addRegAction.triggered.connect(self.on_pbAddReg_clicked)
        elif tableName == "Register":
            self.treeViewPopMenu.addAction(addRegAction)
            addRegAction.triggered.connect(self.on_pbAddReg_clicked)
            self.treeViewPopMenu.addAction(addBfAction)
            addBfAction.triggered.connect(self.on_pbAddBf_clicked)
        elif tableName == "Bitfield":
            self.treeViewPopMenu.addAction(addBfAction)
            addBfAction.triggered.connect(self.on_pbAddBf_clicked)
            self.treeViewPopMenu.addAction(addBfEnumAction)
            addBfEnumAction.triggered.connect(self.on_pbAddBfEnum_clicked)
        elif tableName == "BitfieldEnum":
            self.treeViewPopMenu.addAction(addBfEnumAction)
            addBfEnumAction.triggered.connect(self.on_pbAddBfEnum_clicked)
        
        parent = self.treeViewTableModel.itemFromIndex(index.parent()) if tableName != "MemoryMap" else self.treeViewTableModel.invisibleRootItem()
        if tableName != "MemoryMap" or parent.rowCount() > 1:
            self.treeViewPopMenu.addSeparator()
            delAction = QAction("Delete", self)
            self.treeViewPopMenu.addAction(delAction)
            delAction.triggered.connect(self.do_delete_triggered)
        
        menuPosition = self.ui.treeView.viewport().mapToGlobal(point)
        self.treeViewPopMenu.move(menuPosition)
        self.treeViewPopMenu.show()
    
    @Slot()
    def do_tableView_dataChanged(self, topLeft, bottomRight, roles):
        current = self.ui.treeView.selectedIndexes().pop()
        tableName = str(current.data(RegisterConst.NameRole))
        
        model = None
        idRole = None
        if tableName == "RegisterMap":
            model = self.regMapTableModel
            idRole = RegisterConst.RegMapIdRole
        elif tableName == "Register":
            model = self.regTableModel
            idRole = RegisterConst.RegIdRole
        elif tableName == "Bitfield":
            model = self.bfTableModel
            idRole = RegisterConst.BfIdRole
        elif tableName == "BitfieldEnum":
            model = self.bfEnumTableModel
            idRole = RegisterConst.BfEnumIdRole
        
        # update treeview node accordingly
        # update tree node name if "Name" is changed, or change tree node background color if "Exist" is changed
        if model != None:
            fieldName = model.record().fieldName(bottomRight.column())
            if fieldName == "Name":
                itemId = model.data(model.index(bottomRight.row(), 0), Qt.DisplayRole)
                parentItem = self.treeViewTableModel.itemFromIndex(current.parent())
                for i in range(parentItem.rowCount()):
                    child = current.sibling(i, 0)
                    childId = child.data(idRole)
                    if childId != None and int(childId) == itemId:
                        newName = model.record(bottomRight.row()).value("Name")
                        self.treeViewTableModel.itemFromIndex(child).setData(newName, Qt.DisplayRole)
            elif fieldName == "Exist":
                itemId = model.data(model.index(bottomRight.row(), 0), Qt.DisplayRole) # 'id' is column 0
                parentItem = self.treeViewTableModel.itemFromIndex(current.parent())
                for i in range(parentItem.rowCount()):
                    child = current.sibling(i, 0)
                    childId = child.data(idRole)
                    if childId != None and int(childId) == itemId:
                        if RegisterConst.recordExist(model.record(bottomRight.row())) == False:
                            self.treeViewTableModel.itemFromIndex(child).setData(QColor('grey'), Qt.BackgroundColorRole)
                        else:
                            self.treeViewTableModel.itemFromIndex(child).setData(None, Qt.BackgroundColorRole)

        # resize columns                            
        if tableName == "Register":
            self.ui.tableViewReg.resizeColumnsToContents()            
        else:
            self.ui.tableView.resizeColumnsToContents()
            
        return
    
    @Slot()
    def do_treeView_currentChanged(self, current, previous):
        if self.view == RegisterConst.DesignView:
            tableName = str(current.data(RegisterConst.NameRole))
            if tableName == "MemoryMap": # mem map selected, show mem map table
                self.ui.tableView.setVisible(True)
                self.ui.tableViewReg.setVisible(False)    
                memMapId = int(current.data(RegisterConst.MemMapIdRole))
                self.regMapTableModel.setParentId(memMapId)
                self.regMapTableModel.setFilter("MemoryMapId=%s"%memMapId)
                self.regMapTableModel.select()
                if self.ui.tableView.model() != self.memMapTableModel:
                    self.ui.tableView.setModel(self.memMapTableModel)
                    self.ui.tableView.selectionModel().selectionChanged.connect(self.do_tableView_selectionChanged)
                    if self.__regMapTypeIndex != None:
                        self.ui.tableView.showColumn(self.__regMapTypeIndex)
                        self.__regMapTypeIndex = None
                    if self.__bfValueIndex != None:
                        self.ui.tableView.showColumn(self.__bfValueIndex)
                        self.__bfValueIndex = None
                    self.ui.tableView.hideColumn(0) # id
                    self.ui.tableView.hideColumn(1) # order
                    self.ui.tableView.showColumn(2) # offset address
                    self.ui.tableView.showColumn(3) # name
                    self.ui.tableView.resizeColumnsToContents()
                    
                # update tips
                self.ui.pbAddMemMap.setEnabled(True)
                self.ui.pbAddRegMap.setEnabled(True)
                self.ui.pbAddReg.setEnabled(False)
                self.ui.pbAddBf.setEnabled(False)
                self.ui.pbAddBfEnum.setEnabled(False)
                self.ui.labelDescription.setText("Tips: <br><br>Click <font color=\"red\">%s</font> to add new register map.</br></br>"%self.ui.pbAddRegMap.text())
                
            elif tableName == "RegisterMap": # regmap selected, show regmap table
                self.ui.tableView.setVisible(True)
                self.ui.tableViewReg.setVisible(False)
                memMapId = int(current.data(RegisterConst.MemMapIdRole))
                regMapId = int(current.data(RegisterConst.RegMapIdRole))
                self.regTableModel.setParentId(regMapId)
                self.regTableModel.setFilter("RegisterMapId=%s"%regMapId)
                self.regTableModel.select()
                if self.ui.tableView.model() != self.regMapTableModel or memMapId != self.regMapTableModel.parentId:
                    self.regMapTableModel.setParentId(memMapId)
                    self.regMapTableModel.setFilter("MemoryMapId=%s"%memMapId)
                    self.regMapTableModel.select()
                    self.ui.tableView.setModel(self.regMapTableModel)
                    self.ui.tableView.selectionModel().selectionChanged.connect(self.do_tableView_selectionChanged)
                    if self.__bfValueIndex != None:
                        self.ui.tableView.showColumn(self.__bfValueIndex)
                        self.__bfValueIndex = None
                    self.ui.tableView.hideColumn(0) # id
                    self.ui.tableView.hideColumn(1) # memmap id
                    self.ui.tableView.hideColumn(2) # order
                    self.__regMapTypeIndex = self.regMapTableModel.record().indexOf("Type")
                    self.ui.tableView.hideColumn(self.__regMapTypeIndex)
                    self.ui.tableView.resizeColumnsToContents()
                    self.ui.tableView.resizeColumnsToContents()

                # update tips
                self.ui.pbAddMemMap.setEnabled(False)
                self.ui.pbAddRegMap.setEnabled(True)
                regMapType = current.data(RegisterConst.RegMapTypeRole)
                regMapType = RegisterConst.RegMap if regMapType == None or regMapType == '' else int(regMapType) # default as regmap if not set
                self.ui.pbAddReg.setEnabled(regMapType == RegisterConst.RegMap)                
                self.ui.pbAddBf.setEnabled(False)
                self.ui.pbAddBfEnum.setEnabled(False)
                if regMapType == RegisterConst.RegMap:
                    self.ui.labelDescription.setText("Tips: <br><br>Click <font color=\"red\">%s</font> to add register.</br></br>"%(self.ui.pbAddReg.text()))
                else:
                    self.ui.labelDescription.setText("Tips: <br><br>Assign a filename to sub-module.</br></br>")

            elif tableName == "Register": # reg selected, show reg table
                self.ui.tableView.setVisible(False)
                self.ui.tableViewReg.setVisible(True)                
                regMapId = int(current.data(RegisterConst.RegMapIdRole))
                regId = int(current.data(RegisterConst.RegIdRole))
                self.bfTableModel.setParentId(regId)
                self.bfTableModel.setFilter("RegisterId=%s"%regId)
                self.bfTableModel.select()
                if self.ui.tableViewReg.model() != self.regTableModel or regMapId != self.regTableModel.parentId or self.__treeViewCurrentTable != "Register":
                    self.__regValueIndex = self.regTableModel.record().indexOf("Value")
                    self.regTableModel.setHeaderData(self.__regValueIndex, Qt.Horizontal, "Bits")
                    self.regTableModel.setParentId(regMapId)
                    self.regTableModel.setFilter("RegisterMapId=%s"%regMapId)
                    self.regTableModel.select()
                    self.ui.tableViewReg.setItemDelegateForColumn(self.__regValueIndex, QRegValueDisplayDelegate())
                    if self.ui.tableViewReg.model() != self.regTableModel:
                        self.ui.tableViewReg.setModel(self.regTableModel)
                        self.ui.tableViewReg.selectionModel().selectionChanged.connect(self.do_tableView_selectionChanged)
                    self.ui.tableViewReg.hideColumn(0) # id
                    self.ui.tableViewReg.hideColumn(1) # regmap id
                    self.ui.tableViewReg.hideColumn(2) # order
                    self.ui.tableViewReg.resizeColumnsToContents()

                # update tips
                self.ui.pbAddMemMap.setEnabled(False)
                self.ui.pbAddRegMap.setEnabled(False)
                self.ui.pbAddReg.setEnabled(True)
                self.ui.pbAddBf.setEnabled(True)
                self.ui.pbAddBfEnum.setEnabled(False)                
                self.ui.labelDescription.setText("Tips: <br><br>Click <font color=\"red\">%s</font> to add bitfield</br></br>"%(self.ui.pbAddBf.text()))
                
            elif tableName == "Bitfield": # bf selected, show bf table
                self.ui.tableView.setVisible(True)
                self.ui.tableViewReg.setVisible(False)
                regId = int(current.data(RegisterConst.RegIdRole))
                bfId  = int(current.data(RegisterConst.BfIdRole))
                self.bfEnumTableModel.setParentId(bfId)
                self.bfEnumTableModel.setFilter("BitfieldId=%s"%bfId)
                self.bfEnumTableModel.select()
                if self.ui.tableView.model() != self.bfTableModel or regId != self.bfTableModel.parentId:
                    self.bfTableModel.setParentId(regId)
                    self.bfTableModel.setFilter("RegisterId=%s"%regId)
                    self.bfTableModel.select()
                    self.ui.tableView.setModel(self.bfTableModel)
                    self.ui.tableView.selectionModel().selectionChanged.connect(self.do_tableView_selectionChanged)
                    if self.__regMapTypeIndex != None:
                        self.ui.tableView.showColumn(self.__regMapTypeIndex)
                        self.__regMapTypeIndex = None
                    self.ui.tableView.hideColumn(0) # id
                    self.ui.tableView.hideColumn(1) # regid
                    self.ui.tableView.hideColumn(2) # order
                    self.__bfValueIndex = self.bfTableModel.record().indexOf("Value")
                    self.ui.tableView.hideColumn(self.__bfValueIndex)
                    self.ui.tableView.resizeColumnsToContents()
                    
                # update tips
                self.ui.pbAddMemMap.setEnabled(False)
                self.ui.pbAddRegMap.setEnabled(False)
                self.ui.pbAddReg.setEnabled(False)
                self.ui.pbAddBf.setEnabled(True)
                self.ui.pbAddBfEnum.setEnabled(True)                
                regQ = QSqlQuery("SELECT Width FROM Register WHERE id=%s"%(regId), self.conn)
                text = ""
                while regQ.next(): # only 1 item
                    regW = regQ.value(0)
                    text = "Tips: <pre>"
                    text += RegisterConst.genColoredRegBitsUsage(self.conn, bfId, regId, regW, 13)
                    text += "</pre>"
                self.ui.labelDescription.setText(text)

            elif tableName == "BitfieldEnum": # bfenum selected, show bfenum table
                self.ui.tableView.setVisible(True)
                self.ui.tableViewReg.setVisible(False)                 
                bfId = int(current.data(RegisterConst.BfIdRole))
                if self.ui.tableView.model() != self.bfEnumTableModel or bfId != self.bfEnumTableModel.parentId:
                    self.bfEnumTableModel.setParentId(bfId)
                    self.bfEnumTableModel.setFilter("BitfieldId=%s"%bfId)
                    self.bfEnumTableModel.select()
                    self.ui.tableView.setModel(self.bfEnumTableModel)
                    self.ui.tableView.selectionModel().selectionChanged.connect(self.do_tableView_selectionChanged)
                    if self.__regMapTypeIndex != None:
                        self.ui.tableView.showColumn(self.__regMapTypeIndex)
                        self.__regMapTypeIndex = None
                    if self.__bfValueIndex != None:
                        self.ui.tableView.showColumn(self.__bfValueIndex)
                        self.__bfValueIndex = None
                    self.ui.tableView.hideColumn(0) # id
                    self.ui.tableView.hideColumn(1) # bfid
                    self.ui.tableView.hideColumn(2) # order
                    self.ui.tableView.resizeColumnsToContents()
                    
                # update tips
                self.ui.pbAddMemMap.setEnabled(False)
                self.ui.pbAddRegMap.setEnabled(False)
                self.ui.pbAddReg.setEnabled(False)
                self.ui.pbAddBf.setEnabled(False)
                self.ui.pbAddBfEnum.setEnabled(True)                    
                self.ui.labelDescription.setText("Tips: <br><br>Click <font color=\"red\">%s</font> to add new bitfield enum.</br></br>"%self.ui.pbAddBfEnum.text())

            # select tableView row
            self.__treeViewCurrentTable = tableName
            self.__treeViewCurrentRow = current.row()
            if tableName == "Register":
                tableViewCurrents = self.ui.tableViewReg.selectionModel().selectedIndexes()                
            else:
                tableViewCurrents = self.ui.tableView.selectionModel().selectedIndexes()
            tableViewCurrent = None if len(tableViewCurrents) == 0 else tableViewCurrents[0]
            if tableViewCurrent == None or tableViewCurrent.row() != current.row():
                if tableName == "Register":
                    self.ui.tableViewReg.selectRow(current.row())
                else:
                    self.ui.tableView.selectRow(current.row())                    
        else: # debug view
            tableName = str(current.data(RegisterConst.NameRole))
            if tableName != "MemoryMap":          
                for regDebugModel in self.regDebugModels:
                    regMapId = int(current.data(RegisterConst.RegMapIdRole))
                    if regDebugModel.id == regMapId:
                        self.ui.tableView.setModel(regDebugModel)
                        self.ui.tableView.selectionModel().selectionChanged.connect(self.do_tableView_selectionChanged)
                        if self.__regMapTypeIndex != None:
                            self.ui.tableView.showColumn(self.__regMapTypeIndex)
                            self.__regMapTypeIndex = None
                        self.ui.tableView.showColumn(0)
                        self.ui.tableView.showColumn(1)
                        self.ui.tableView.showColumn(2)
                        self.ui.tableView.resizeColumnsToContents()
                        break
        return

    @Slot()
    def on_pbAddMemMap_clicked(self):
        current = self.ui.treeView.selectedIndexes().pop()
        tableName = str(current.data(RegisterConst.NameRole))
        newMemMapRowIndex = current.row() + 1
        r = self.newMemMapRow(self.memMapTableModel, newMemMapRowIndex)
        
        newMemMapItem = QStandardItem(self.moduleIcon, r.value("name"))
        newMemMapItem.setData("MemoryMap", RegisterConst.NameRole)
        newMemMapItem.setData(r.value("id"), RegisterConst.MemMapIdRole)

        root = self.treeViewTableModel.invisibleRootItem()
        if tableName == "MemoryMap":
            if (current.row() + 1) == root.rowCount(): # current is last one
                root.appendRow(newMemMapItem)
                item = root.child(current.row() + 1)
                self.ui.treeView.selectionModel().setCurrentIndex(item.index(), QItemSelectionModel.ClearAndSelect)
            else:
                root.insertRows(current.row() + 1, 1)
                root.setChild(current.row() + 1, newMemMapItem)
                item = root.child(current.row() + 1)
                self.ui.treeView.selectionModel().setCurrentIndex(item.index(), QItemSelectionModel.ClearAndSelect)       
        return

    @Slot()
    def on_pbAddRegMap_clicked(self):
        current = self.ui.treeView.selectedIndexes().pop()
        tableName = str(current.data(RegisterConst.NameRole))
        memoryMapId = int(current.data(RegisterConst.MemMapIdRole))
        newRegMapRowIndex = -1 if tableName != "RegisterMap" else current.row() + 1
        r = self.newRegMapRow(self.regMapTableModel, memoryMapId, newRegMapRowIndex)
        
        newRegMapItem = QStandardItem(self.regMapIcon, r.value("name"))
        newRegMapItem.setData("RegisterMap", RegisterConst.NameRole)
        newRegMapItem.setData(memoryMapId, RegisterConst.MemMapIdRole)
        newRegMapItem.setData(r.value("id"), RegisterConst.RegMapIdRole)
        newRegMapItem.setData(RegisterConst.RegMap, RegisterConst.RegMapTypeRole)
        
        standardItem = self.treeViewTableModel.itemFromIndex(current)
        if tableName == "MemoryMap":
            standardItem.appendRow(newRegMapItem)
        elif tableName == "RegisterMap":
            if (current.row() + 1) == standardItem.parent().rowCount(): # current is last one
                standardItem.parent().appendRow(newRegMapItem)
                item = standardItem.parent().child(current.row() + 1)
                self.ui.treeView.selectionModel().setCurrentIndex(item.index(), QItemSelectionModel.ClearAndSelect)
            else:
                standardItem.parent().insertRows(current.row() + 1, 1)
                standardItem.parent().setChild(current.row() + 1, newRegMapItem)
                item = standardItem.parent().child(current.row() + 1)
                self.ui.treeView.selectionModel().setCurrentIndex(item.index(), QItemSelectionModel.ClearAndSelect) 
        return

    @Slot()
    def do_pbAddRegMod_clicked(self):
        current = self.ui.treeView.selectedIndexes().pop()
        tableName = str(current.data(RegisterConst.NameRole))
        memoryMapId = int(current.data(RegisterConst.MemMapIdRole))
        newRegMapRowIndex = -1 if tableName != "RegisterMap" else current.row() + 1
        r = self.newRegMapRow(self.regMapTableModel, memoryMapId, newRegMapRowIndex, RegisterConst.RegMod)
        
        newRegMapItem = QStandardItem(self.regMapIcon, r.value("name"))
        newRegMapItem.setData("RegisterMap", RegisterConst.NameRole)
        newRegMapItem.setData(memoryMapId, RegisterConst.MemMapIdRole)
        newRegMapItem.setData(r.value("id"), RegisterConst.RegMapIdRole)
        newRegMapItem.setData(RegisterConst.RegMod, RegisterConst.RegMapTypeRole)
        
        standardItem = self.treeViewTableModel.itemFromIndex(current)
        if tableName == "MemoryMap":
            standardItem.appendRow(newRegMapItem)
        elif tableName == "RegisterMap":
            if (current.row() + 1) == standardItem.parent().rowCount(): # current is last one
                standardItem.parent().appendRow(newRegMapItem)
                item = standardItem.parent().child(current.row() + 1)
                self.ui.treeView.selectionModel().setCurrentIndex(item.index(), QItemSelectionModel.ClearAndSelect)
            else:
                standardItem.parent().insertRows(current.row() + 1, 1)
                standardItem.parent().setChild(current.row() + 1, newRegMapItem)
                item = standardItem.parent().child(current.row() + 1)
                self.ui.treeView.selectionModel().setCurrentIndex(item.index(), QItemSelectionModel.ClearAndSelect)     
        return

    @Slot()
    def on_pbAddReg_clicked(self):
        current = self.ui.treeView.selectedIndexes().pop()
        tableName = str(current.data(RegisterConst.NameRole))
        memoryMapId = int(current.data(RegisterConst.MemMapIdRole))
        regMapId = int(current.data(RegisterConst.RegMapIdRole))
        newRegRowIndex = -1 if tableName != "Register" else current.row() + 1
        r = self.newRegRow(self.regTableModel, regMapId, 0, 32, newRegRowIndex)
        
        newRegItem = QStandardItem(self.regIcon, r.value("name"))
        newRegItem.setData("Register", RegisterConst.NameRole)
        newRegItem.setData(memoryMapId, RegisterConst.MemMapIdRole)
        newRegItem.setData(regMapId, RegisterConst.RegMapIdRole)
        newRegItem.setData(r.value("id"), RegisterConst.RegIdRole)
    
        standardItem = self.treeViewTableModel.itemFromIndex(current)
        if tableName == "RegisterMap":
            standardItem.appendRow(newRegItem)
        elif tableName == "Register":
            if (current.row() + 1) == standardItem.parent().rowCount(): # current is last one
                standardItem.parent().appendRow(newRegItem)
                item = standardItem.parent().child(current.row() + 1)
                self.ui.treeView.selectionModel().setCurrentIndex(item.index(), QItemSelectionModel.ClearAndSelect)
            else:
                standardItem.parent().insertRows(current.row() + 1, 1)
                standardItem.parent().setChild(current.row() + 1, newRegItem)
                item = standardItem.parent().child(current.row() + 1)
                self.ui.treeView.selectionModel().setCurrentIndex(item.index(), QItemSelectionModel.ClearAndSelect)          
        return

    @Slot()
    def on_pbAddBf_clicked(self):
        current = self.ui.treeView.selectedIndexes().pop()
        tableName = str(current.data(RegisterConst.NameRole))
        memoryMapId = int(current.data(RegisterConst.MemMapIdRole))
        regMapId = int(current.data(RegisterConst.RegMapIdRole))
        regId = int(current.data(RegisterConst.RegIdRole))
        newBfRowIndex = -1 if tableName != "Bitfield" else current.row() + 1 
        r = self.newBfRow(self.bfTableModel, regId, 8, newBfRowIndex)
        
        newBfItem = QStandardItem(self.bfIcon, r.value("name"))
        newBfItem.setData("Bitfield", RegisterConst.NameRole)
        newBfItem.setData(memoryMapId, RegisterConst.MemMapIdRole)
        newBfItem.setData(regMapId, RegisterConst.RegMapIdRole)
        newBfItem.setData(regId, RegisterConst.RegIdRole)
        newBfItem.setData(r.value("id"), RegisterConst.BfIdRole)

        standardItem = self.treeViewTableModel.itemFromIndex(current)
        if tableName == "Register":
            standardItem.appendRow(newBfItem)
        elif tableName == "Bitfield":
            if (current.row() + 1) == standardItem.parent().rowCount(): # current is last one
                standardItem.parent().appendRow(newBfItem)
                item = standardItem.parent().child(current.row() + 1)
                self.ui.treeView.selectionModel().setCurrentIndex(item.index(), QItemSelectionModel.ClearAndSelect)
            else:
                standardItem.parent().insertRows(current.row() + 1, 1)
                standardItem.parent().setChild(current.row() + 1, newBfItem)
                item = standardItem.parent().child(current.row() + 1)
                self.ui.treeView.selectionModel().setCurrentIndex(item.index(), QItemSelectionModel.ClearAndSelect)
        return

    @Slot()
    def on_pbAddBfEnum_clicked(self):
        current = self.ui.treeView.selectedIndexes().pop()
        tableName = str(current.data(RegisterConst.NameRole))
        memoryMapId = int(current.data(RegisterConst.MemMapIdRole))
        regMapId = int(current.data(RegisterConst.RegMapIdRole))
        regId = int(current.data(RegisterConst.RegIdRole))
        bfId = int(current.data(RegisterConst.BfIdRole))
        newBfEnumRowIndex = -1 if tableName != "BitfieldEnum" else current.row() + 1
        r = self.newBfEnumRow(self.bfEnumTableModel, bfId, newBfEnumRowIndex)
        
        newBfEnumItem = QStandardItem(self.bfenumIcon, r.value("name"))
        newBfEnumItem.setData("BitfieldEnum", RegisterConst.NameRole)
        newBfEnumItem.setData(memoryMapId, RegisterConst.MemMapIdRole)
        newBfEnumItem.setData(regMapId, RegisterConst.RegMapIdRole)
        newBfEnumItem.setData(regId, RegisterConst.RegIdRole)
        newBfEnumItem.setData(bfId, RegisterConst.BfIdRole)
        newBfEnumItem.setData(r.value("id"), RegisterConst.BfEnumIdRole)
        
        standardItem = self.treeViewTableModel.itemFromIndex(current)
        if tableName == "Bitfield":
            standardItem.appendRow(newBfEnumItem)
        elif tableName == "BitfieldEnum":
            if (current.row() + 1) == standardItem.parent().rowCount(): # current is last one
                standardItem.parent().appendRow(newBfEnumItem)
                item = standardItem.parent().child(current.row() + 1)
                self.ui.treeView.selectionModel().setCurrentIndex(item.index(), QItemSelectionModel.ClearAndSelect)
            else:
                standardItem.parent().insertRows(current.row() + 1, 1)
                standardItem.parent().setChild(current.row() + 1, newBfEnumItem)
                item = standardItem.parent().child(current.row() + 1)
                self.ui.treeView.selectionModel().setCurrentIndex(item.index(), QItemSelectionModel.ClearAndSelect)
        return
    
    @Slot()
    def do_delete_triggered(self):
        current = self.ui.treeView.selectedIndexes().pop()
        tableName = str(current.data(RegisterConst.NameRole))

        # remove node
        parent = self.treeViewTableModel.itemFromIndex(current.parent()) if tableName != "MemoryMap" else self.treeViewTableModel.invisibleRootItem()
        if tableName != "MemoryMap" or parent.rowCount() > 1:
            # remove from table
            if tableName == "MemoryMap":
                memMapId = int(current.data(RegisterConst.MemMapIdRole))
                query = QSqlQuery(self.conn)
                query.exec_("DELETE FROM MemoryMap WHERE id=%s"%memMapId)
                self.memMapTableModel.select()
                self.regMapTableModel.select()
                self.regTableModel.select()
                self.bfTableModel.select()
                self.bfEnumTableModel.select()
            elif tableName == "RegisterMap":
                regMapId = int(current.data(RegisterConst.RegMapIdRole))
                query = QSqlQuery(self.conn)
                query.exec_("DELETE FROM RegisterMap WHERE id=%s"%regMapId)
                self.regMapTableModel.select()
                self.regTableModel.select()
                self.bfTableModel.select()
                self.bfEnumTableModel.select()
            elif tableName == "Register":
                regId = int(current.data(RegisterConst.RegIdRole))
                query = QSqlQuery(self.conn)
                query.exec_("DELETE FROM Register WHERE id=%s"%regId)
                self.regTableModel.select()
                self.bfTableModel.select()
                self.bfEnumTableModel.select()
            elif tableName == "Bitfield":
                bfId = int(current.data(RegisterConst.BfIdRole))
                query = QSqlQuery(self.conn)
                query.exec_("DELETE FROM Bitfield WHERE id=%s"%bfId)
                self.bfTableModel.select()
                self.bfEnumTableModel.select()
            elif tableName == "BitfieldEnum":
                bfEnumId = int(current.data(RegisterConst.BfEnumIdRole))
                query = QSqlQuery(self.conn)
                query.exec_("DELETE FROM BitfieldEnum WHERE id=%s"%bfEnumId)
                self.bfEnumTableModel.select()
            # remove from tree
            parent.removeRow(current.row())
        
        return