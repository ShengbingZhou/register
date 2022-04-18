# built-in package
import os
import sys
import re

# pyside2 package
from PySide2.QtWidgets import QFileDialog, QMessageBox, QProgressDialog
from PySide2.QtCore import Qt, QDir, QCoreApplication
from PySide2.QtSql import QSqlQuery, QSqlQueryModel
from PySide2.QtGui import QColor

# python-docx package
from docx import Document, oxml, shared
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches

class QRegisterConst:

    # tool version
    Version = "0.1.0-(beta)"
    
    # Base Directory
    BaseDir = getattr(sys, '_MEIPASS', os.path.abspath(os.path.dirname(__file__)))

    # style file
    StyleFile = os.path.join(BaseDir, "style/style.qss")

    # design file ext
    DesignFileExt = ".reg"
    
    # reg access log file ext
    RegLogFileExt = ".reglog"
    
    # tab type const
    WelcomeTab = 0
    ModuleTab  = 1
    RegLogTab  = 2
    
    # module view const
    DesignView  = 0
    DebugView   = 1

    # regmap type
    RegMap = 0
    RegMod = 1

    # bf access options
    AccessTypes = ['read-write', 'read-only', 'write-only', 'read-writeOnce', 'writeOnce']

    # bf reset types
    ResetTypes = ['async', 'sync']

    # visibility options
    VisibilityOptions = ['public', 'private']

    # bool options
    BoolOptions = ['true', 'false']

    # special role for treeview item
    TableNameRole  = Qt.UserRole + 1
    infoIdRole     = Qt.UserRole + 2
    MemMapIdRole   = Qt.UserRole + 3
    RegMapIdRole   = Qt.UserRole + 4
    RegIdRole      = Qt.UserRole + 5
    BfIdRole       = Qt.UserRole + 6
    BfEnumIdRole   = Qt.UserRole + 7
    RegMapTypeRole = Qt.UserRole + 8
    BfReadOnlyRole = Qt.UserRole + 9

    # hardware access driver
    RegisterAccessDriverClass= None
    
    # value column index in debug view
    ValueColumnOfDebugView = 3

    @staticmethod
    def strToInt(text):
        if text is None:
            return 0
        if text.startswith("0x"):         # 0x1234
            return int(text, 16)
        if text.startswith("'h"):         # 'h1234
            text = text.replace("'h", "")
            return int(text, 16)
        if text.startswith("'d"):
            text = text.replace("'d", "") # 'd1234
            return int(text)
        if text.startswith("'b"):
            text = text.replace("'b", "") # 'b1011
            return int(text, 2)
        if "'h" in text:                  # 16'h1234
            t = text.split("'h")
            return int(t[1], 16)
        if "'d" in text:                  # 16'd1234
            t = text.split("'d")
            return int(t[1])
        if "'b" in text:                  # 16'b1011
            t = text.split("'b")
            return int(t[1], 2)
        return int(text)

    @staticmethod
    def recordExist(record):
        exist = record.value("Exist")
        if exist is None:
            return True
        else:
            if exist == 'false' or exist == '0' or exist == "n" or exist == 'no':
                return False
            else:
                return True

    @staticmethod
    def isReadOnly(value):
        if value is None:
            return False
        valueString = str(value)
        if valueString is '':
            return False
        if 'w' in valueString:
            return False
        else:
            return True

    @staticmethod
    def findRegAccessDriverClass():
        if QRegisterConst.RegisterAccessDriverClass is None:
            if os.path.isfile(QDir.homePath() + "/QRegisterAccessDriver/QRegisterAccess.py"):        
                DriverPath  = QDir.homePath() + "/QRegisterAccessDriver"
                sys.path.append(DriverPath)
                DriverMod   = __import__("QRegisterAccess")
                QRegisterConst.RegisterAccessDriverClass = getattr(DriverMod, "QRegisterAccess")        

    @staticmethod
    def genColoredRegBitsUsage(conn, bfId, regId, regW, fontSize):
        if regW == None:
            return

        bfColorsIndex = 0
        bfColors  = ["DarkSeaGreen",  "LightSalmon",     "PowderBlue",     "LightPink",      "Aquamarine",     "Bisque",         "LightSteelBlue", "DarkKhaki"]         
        bfQColors = [QColor(0x8FBC8F), QColor(0xFFA07A), QColor(0xB0E0E6), QColor(0xFFB6C1), QColor(0x66CDAA), QColor(0xFFE4C4), QColor(0xB0C4DE), QColor(0xBDB76B)]
        value = []

        regW = int(regW)
        regB = regW - 1
        bfQuery = QSqlQuery("SELECT * FROM Bitfield WHERE RegisterId=%s ORDER BY CAST(RegisterOffset as int) DESC"%(regId), conn)
        while bfQuery.next():
            _bfId   = bfQuery.value("id")
            _regOff = QRegisterConst.strToInt(bfQuery.value("RegisterOffset"))
            _bfW    = int(bfQuery.value("Width"))

            # unused bits before bitfield 
            if _bfW > 0 and regB > (_regOff + _bfW - 1):
                text  = ""
                start = _regOff + _bfW
                end   = regB + 1
                for i in range(start, end):
                    text += "%s,"%(regB) if i < (end - 1) and regB > 0 else "%s"%(regB)
                    regB -= 1
                    if regB < 0:
                        break
                value.append((None, text))

            # bitfield bits
            if _bfW > 0 and regB >= 0:
                text = ""
                start = _regOff
                end   = _regOff + _bfW               
                for j in range(_regOff, _regOff + _bfW):                    
                    text += "%s,"%(regB) if j < (end - 1) and regB > 0 else "%s"%(regB)
                    regB -= 1
                    if regB < 0:
                        break
                if bfId == _bfId:
                    value.append((bfQColors[bfColorsIndex], text, _bfId, 1))
                else:
                    value.append((bfQColors[bfColorsIndex], text, _bfId))
                bfColorsIndex = 0 if (bfColorsIndex + 1) >= len(bfQColors) else bfColorsIndex + 1

        # left unsed bits
        if regB >= 0:
            text = ""
            for k in range(0, regB + 1):
                text += "%s,"%(regB) if regB > 0 else "%s"%(regB)
                regB -= 1
            value.append((None, text))

        return value

    @staticmethod
    def genRegValueFromBitfields(conn, regId):
        regValue = 0
        bfQuery = conn.exec_("SELECT * FROM Bitfield WHERE RegisterId=%s"%(regId))
        while bfQuery.next():
            regOff    = QRegisterConst.strToInt(bfQuery.value("RegisterOffset"))
            bfDefault = QRegisterConst.strToInt(bfQuery.value("DefaultValue"))
            regValue += bfDefault << regOff
        return regValue

    @staticmethod
    def exporDocx(parent, conn):
        fileName, filterUsed = QFileDialog.getSaveFileName(parent, "Export Word file", QDir.homePath(), "Word File (*.docx)")
        if fileName == '':           
            return

        f_name, f_ext = os.path.splitext(os.path.basename(fileName))
        if f_ext != ".docx":
            fileName += ".docx"
        docx = Document()
        docx.styles['Heading 1'].font.size = shared.Pt(11)
        docx.styles['Heading 2'].font.size = shared.Pt(10)
        docx.styles['Heading 3'].font.size = shared.Pt(9)
        docx.styles['Heading 4'].font.size = shared.Pt(8)
        docx.styles['Normal'].font.size    = shared.Pt(8)
                    
        # memory map
        memoryMapQueryModel = QSqlQueryModel()
        memoryMapQueryModel.setQuery("SELECT * FROM MemoryMap", conn)

        title = docx.add_heading('MemoryMap Table\n', level = 1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fields = ['Name', 'Description']
        table = docx.add_table(rows=memoryMapQueryModel.rowCount() + 1, cols=len(fields), style='Table Grid')

        for i, row in enumerate(table.rows):
            for j, (cell, field) in enumerate(zip(row.cells, fields)):
                if i == 0: # table header
                    cell.text = fields[j]
                    cell._tc.get_or_add_tcPr().append(oxml.parse_xml(r'<w:shd {} w:fill="c0c0c0"/>'.format(oxml.ns.nsdecls('w'))))
                else:
                    memMapRecord = memoryMapQueryModel.record(i - 1)
                    if field == 'Name':
                        cell.text = memMapRecord.value("Name")
        docx.add_page_break()

        # setup progress dialog
        dlgProgress = QProgressDialog(parent)
        dlgProgress.setWindowTitle("Exporting ...")
        dlgProgress.setWindowModality(Qt.WindowModal)
        dlgProgress.setMinimum(0)
        dlgProgress.show()

        for i in range(memoryMapQueryModel.rowCount()):
            memMapRecord = memoryMapQueryModel.record(i)
            docx.add_heading('MemoryMap: %s'%(memMapRecord.value("Name")), level = 2)

            # register map
            regMapQueryModel = QSqlQueryModel()
            regMapQueryModel.setQuery("SELECT * FROM RegisterMap WHERE memoryMapId=%s ORDER BY DisplayOrder ASC"%memMapRecord.value("id"), conn)
            
            # progress dialog
            dlgProgress.setMaximum(memoryMapQueryModel.rowCount())
            dlgProgress.setValue(0)
            QCoreApplication.processEvents()              
            
            for j in range(regMapQueryModel.rowCount()):
                regMapRecord = regMapQueryModel.record(j)
                docx.add_heading('RegisterMap: %s'%(regMapRecord.value("Name")), level = 3)
                docx.add_paragraph("Description : %s\n" \
                                "BaseAddress : %s"%(regMapRecord.value("Description"), regMapRecord.value("OffsetAddress")))

                # update progress dialog
                dlgProgress.setLabelText("Exporting register map '%s' to %s "%(regMapRecord.value("Name"), fileName))
                dlgProgress.setValue(j)
                QCoreApplication.processEvents()    

                # register
                regQueryModel = QSqlQueryModel()
                regQueryModel.setQuery("SELECT * FROM Register WHERE RegisterMapId=%s ORDER BY DisplayOrder ASC"%regMapRecord.value("id"), conn)

                fields = ['Name', 'Address', 'Description']
                table = docx.add_table(1, cols=len(fields), style='Table Grid')
                for c, (cell, field) in enumerate(zip(table.rows[0].cells, fields)):
                    cell.text = fields[c]
                    cell._tc.get_or_add_tcPr().append(oxml.parse_xml(r'<w:shd {} w:fill="c0c0c0"/>'.format(oxml.ns.nsdecls('w'))))
                regRe = re.compile('\d+:\d+')                
                for r in range(regQueryModel.rowCount()):
                    regRecord = regQueryModel.record(r)
                    regWidth = int(regRecord.value("Width"))
                    regDesc = regRecord.value("Description")
                    regMatch = regRe.match(regRecord.value("Array"))
                    if regMatch is None:
                        regName = regRecord.value("Name")
                        regAddr = "%s"%regRecord.value("OffsetAddress")                        
                        row = table.add_row()
                        row.cells[0].text = regName
                        row.cells[1].text = regAddr
                        row.cells[2].text = regDesc                        
                    else:                        
                        regArray = regMatch.string.split(':')
                        regArray0 = int(regArray[0])
                        regArray1 = int(regArray[1])
                        start = min(regArray0, regArray1)
                        end   = max(regArray0, regArray1)
                        for regI in range(start, end + 1):
                            regName = "%s%s"%(regRecord.value("Name"), regI)
                            regAddr = hex(QRegisterConst.strToInt(regRecord.value("OffsetAddress")) + int(regWidth * (regI - start) / 8))
                            row = table.add_row()
                            row.cells[0].text = regName
                            row.cells[1].text = regAddr
                            row.cells[2].text = regDesc

                for k in range(regQueryModel.rowCount()):
                    regRecord = regQueryModel.record(k)
                    regMatch = regRe.match(regRecord.value("Array"))
                    if regMatch is None:
                        docx.add_heading('Register: %s'%(regRecord.value("Name")), level = 4)
                        docx.add_paragraph('Description : %s\n' \
                                           'Address : %s'%(regRecord.value("Description"), regRecord.value("OffsetAddress")))
                    else:
                        regArray = regMatch.string.split(':')
                        regArray0 = int(regArray[0])
                        regArray1 = int(regArray[1])
                        start = min(regArray0, regArray1)
                        end   = max(regArray0, regArray1)
                        regAddrStart = hex(QRegisterConst.strToInt(regRecord.value("OffsetAddress")))
                        regAddrend   = hex(QRegisterConst.strToInt(regRecord.value("OffsetAddress")) + int(regWidth * (end - start) / 8))
                        docx.add_heading('Register: %s%s ~ %s%s'%(regRecord.value("Name"), start, regRecord.value("Name"), end), level = 4)
                        docx.add_paragraph('Description : %s\n' \
                                           'Address : %s ~ %s'%(regRecord.value("Description"), regAddrStart, regAddrend))

                    # bitfield
                    bfQueryModel = QSqlQueryModel()
                    bfQueryModel.setQuery("SELECT * FROM Bitfield WHERE RegisterId=%s ORDER BY DisplayOrder ASC"%regRecord.value("id"), conn)

                    fields = ['Name', 'Bits', 'ResetValue', 'Description']
                    table = docx.add_table(rows=bfQueryModel.rowCount() + 1, cols=len(fields), style='Table Grid')
                    table.allow_autofit = True
                    for r, row in enumerate(table.rows):
                        for c, (cell, field) in enumerate(zip(row.cells, fields)):
                            if r == 0:
                                cell.text = fields[c]
                                cell._tc.get_or_add_tcPr().append(oxml.parse_xml(r'<w:shd {} w:fill="c0c0c0"/>'.format(oxml.ns.nsdecls('w'))))
                            else:
                                bfRecord = bfQueryModel.record(r - 1)
                                if field == 'Name':
                                    cell.text = bfRecord.value("Name")
                                if field == 'Bits':
                                    cell.text = "[%s:%s]"%(int(bfRecord.value("Width")) + QRegisterConst.strToInt(bfRecord.value("RegisterOffset")) - 1, QRegisterConst.strToInt(bfRecord.value("RegisterOffset")))
                                if field == 'ResetValue':
                                    cell.text = "%s"%(bfRecord.value("DefaultValue"))
                                if field == 'Description':
                                    cell.text = bfRecord.value("Description")
                docx.add_page_break()
        docx.add_page_break()
        docx.save(fileName)
        dlgProgress.close()
        QMessageBox.information(parent, "Exporting docx", "Done!", QMessageBox.Yes)
        return

    @staticmethod
    def exportVerilog(parent, conn):
        folder = QFileDialog.getExistingDirectory(parent, "Export Verilog file", QDir.homePath())
        if os.path.exists(folder) is False:
            return
        if parent.newModule is True:
            f_name, f_ext = os.path.splitext(os.path.basename(parent.newFileName))
        else:
            f_name, f_ext = os.path.splitext(os.path.basename(parent.fileName))
        folder = folder + "/" + f_name
        if os.path.exists(folder) is False:
            os.mkdir(folder)
            
        # info
        infoQueryModel = QSqlQueryModel()
        infoQueryModel.setQuery("SELECT * FROM info", conn)
        infoRecord = infoQueryModel.record(0)
        moduleName = infoRecord.value("Name")

        # memory map
        memoryMapQueryModel = QSqlQueryModel()
        memoryMapQueryModel.setQuery("SELECT * FROM MemoryMap", conn)

        # output uvm top
        svUVMTopFileName = folder + "/" + moduleName.lower() + "_top.sv"
        svUvmTopFile = open(svUVMTopFileName, "w")        
        for i in range(memoryMapQueryModel.rowCount()):
            memMapRecord  = memoryMapQueryModel.record(i)
            memMapName    = memMapRecord.value("Name")
            svUvmTopFile.write("`include \"%s\"\n"%(moduleName.lower() + "_" + memMapName.lower()  + ".sv"))
        svUvmTopFile.write("\n")
        svUvmTopFile.write("class %s extends uvm_reg_block;\n"%(moduleName.lower()))
        for i in range(memoryMapQueryModel.rowCount()):
            memMapRecord = memoryMapQueryModel.record(i)
            memMapName   = memMapRecord.value("Name")
            memMapNameSV = "mm_" + memMapName
            svUvmTopFile.write("    rand %s %s;\n"%(memMapNameSV.lower(), memMapName.lower()))
        svUvmTopFile.write("    virtual function void build();\n")
        svUvmTopFile.write("        default_map = create_map(\"default_map\", 0, 4, UVM_LITTLE_ENDIAN, 1);\n") # TODO: param is not meaningful for top memory map?
        for i in range(memoryMapQueryModel.rowCount()):
            memMapRecord = memoryMapQueryModel.record(i)
            memMapName   = memMapRecord.value("Name")
            memMapNameSV = "mm_" + memMapName
            svUvmTopFile.write("        %s = %s::type_id::create(\"%s\");\n"%(memMapName.lower(), memMapNameSV.lower(), memMapName.lower()))
            svUvmTopFile.write("        %s.configure(this, \"\");\n"%(memMapName.lower()))
            svUvmTopFile.write("        %s.build();\n"%(memMapName.lower()))
            svUvmTopFile.write("        default_map.add_submap(%s.default_map, 0);\n"%(memMapName.lower()))
            svUvmTopFile.write("        %s.lock_model();\n"%(memMapName.lower()))
            svUvmTopFile.write("\n")
        svUvmTopFile.write("    endfunction\n")
        svUvmTopFile.write("    `uvm_object_utils(%s);\n"%(moduleName.lower()))
        svUvmTopFile.write("    function new(input string name = \"%s\");\n"%(moduleName.lower()))
        svUvmTopFile.write("        super.new(name, UVM_NO_COVERAGE);\n")
        svUvmTopFile.write("    endfunction\n")
        svUvmTopFile.write("endclass\n")

        # output uvm module and sv header
        for i in range(memoryMapQueryModel.rowCount()):
            memMapRecord   = memoryMapQueryModel.record(i)
            memMapName     = memMapRecord.value("Name")
            memMapNameSV   = "mm_" + memMapName
            memMapAddr     = memMapRecord.value("OffsetAddress")
            memMapEndian   = memMapRecord.value("Endianness")
            memMapAddrUnit = memMapRecord.value("AddressUnitBits")
            memMapAddrBytesW = 1 if memMapAddrUnit is None or memMapAddrUnit is '' else QRegisterConst.strToIn(memMapAddrUnit) / 8

            # sv header
            svHeaderBaseFileName = moduleName.lower() + "_" + memMapName.lower()  + "_sv.h"
            svHeaderFileName = folder + "/" + svHeaderBaseFileName
            svHeaderFile = open(svHeaderFileName, "w")
            svHeaderLines = []

            # uvm memmap file
            svUVMMemMapBaseFileName = moduleName.lower() + "_" + memMapName.lower()  + ".sv"
            svUVMMemMapFileName = folder + "/" + svUVMMemMapBaseFileName
            svUvmMemMapFile = open(svUVMMemMapFileName, "w")

            # register map
            regMapQueryModel = QSqlQueryModel()
            regMapQueryModel.setQuery("SELECT * FROM RegisterMap WHERE memoryMapId=%s ORDER BY DisplayOrder ASC"%memMapRecord.value("id"), conn)

            # uvm memap block to include all regmaps
            for j in range(regMapQueryModel.rowCount()):
                regMapRecord = regMapQueryModel.record(j)
                regMapName   = regMapRecord.value("Name")
                svUvmRegMapBaseFileName = moduleName.lower() + "_" + memMapName.lower() + "_" + regMapName.lower() + ".sv"
                svUvmMemMapFile.write("`include \"%s\"\n"%(svUvmRegMapBaseFileName))
            svUvmMemMapFile.write("\n")
            
            for j in range(regMapQueryModel.rowCount()):
                regMapRecord = regMapQueryModel.record(j)
                regMapName   = regMapRecord.value("Name")
                regMapNameSV = "rm_%s_%s"%(memMapName, regMapName)
                regMapAddr   = regMapRecord.value("OffsetAddress")

                # register
                regQueryModel = QSqlQueryModel()
                regQueryModel.setQuery("SELECT * FROM Register WHERE RegisterMapId=%s ORDER BY DisplayOrder ASC"%regMapRecord.value("id"), conn)

                # create regmap class, define reg variable in regmap class in memmap file
                svUvmMemMapFile.write("class %s extends uvm_reg_block;\n"%(regMapNameSV.lower()))
                for k in range(regQueryModel.rowCount()):
                    regRecord = regQueryModel.record(k)
                    regName   = regRecord.value("Name")
                    regNameSV = "rg_%s_%s_%s"%(memMapName, regMapName, regName)
                    regRe     = re.compile('\d+:\d+')  
                    regMatch  = regRe.match(regRecord.value("Array"))
                    if regMatch is None:                    
                        svUvmMemMapFile.write("    rand %s %s;\n"%(regNameSV.lower(), regName.lower()))
                    else:
                        regArray = regMatch.string.split(':')
                        regArray0 = int(regArray[0])
                        regArray1 = int(regArray[1])
                        start = min(regArray0, regArray1)
                        end   = max(regArray0, regArray1)
                        for regI in range(start, end + 1):
                            regNameI   = regName   + str(regI)
                            regNameSVI = regNameSV + str(regI)                        
                            svUvmMemMapFile.write("    rand %s %s;\n"%(regNameSVI.lower(), regNameI.lower()))
                svUvmMemMapFile.write("    virutal function void build();\n")
                svUvmMemMapFile.write("        default_map = create_map(\"default_map\", 0, 4, UVM_LITTLE_ENDIAN, 1);\n") # TODO: param is not meaningful for top memory map?

                # create regmap file
                svUVMRegMapBaseFileName = moduleName.lower() + "_" + memMapName.lower() + "_" + regMapName.lower() + ".sv"
                svUVMRegMapFileName = folder + "/" + svUVMRegMapBaseFileName
                svUvmRegMapFile = open(svUVMRegMapFileName, "w")
                svUvmRegMapFile.write("`ifndef __%s__\n"%(svUVMRegMapBaseFileName.replace(".", "_")))
                svUvmRegMapFile.write("`define __%s__\n"%(svUVMRegMapBaseFileName.replace(".", "_")))
                svUvmRegMapFile.write("\n")

                for k in range(regQueryModel.rowCount()):
                    regRecord  = regQueryModel.record(k)
                    regId      = regRecord.value("id")
                    regName    = regRecord.value("Name")
                    regAddr    = regRecord.value("OffsetAddress")
                    regWidth   = QRegisterConst.strToInt(regRecord.value("Width"))
                    regDefault = QRegisterConst.genRegValueFromBitfields(conn, regId)

                    regNameSV = "rg_%s_%s_%s"%(memMapName, regMapName, regName)
                    regAddrSV = QRegisterConst.strToInt(memMapAddr) + QRegisterConst.strToInt(regMapAddr) + QRegisterConst.strToInt(regAddr)

                    # bitfield
                    bfQueryModel = QSqlQueryModel()
                    bfQueryModel.setQuery("SELECT * FROM Bitfield WHERE RegisterId=%s ORDER BY DisplayOrder ASC"%regRecord.value("id"), conn)

                    regRe     = re.compile('\d+:\d+')  
                    regMatch  = regRe.match(regRecord.value("Array"))
                    if regMatch is None:
                        # sv header
                        svHeaderLines.append("`define %s %s\n"%(regNameSV.upper() + "_ADDR", hex(regAddrSV).replace("0x", "'h")))
                        svHeaderLines.append("`define %s %s\n"%(regNameSV.upper() + "_RST",  hex(regDefault).replace("0x", "'h")))

                        # create register variable to regmap class of memmap file
                        svUvmMemMapFile.write("        %s = %s::type_id::create(\"%s\");\n"%(regName.lower(), regNameSV.lower(), regName.lower()))
                        if bfQueryModel.rowCount() == 1:
                            bfRecord  = bfQueryModel.record(0)
                            bfName    = bfRecord.value("Name")
                            bfWidth   = QRegisterConst.strToInt(bfRecord.value("Width"))
                            regOffset = QRegisterConst.strToInt(bfRecord.value("RegisterOffset"))
                            bfDefault = QRegisterConst.strToInt(bfRecord.value("DefaultValue"))
                            svUvmMemMapFile.write("        %s.configure(this, null, \"%s[%s:0]\");\n"%(regName.lower(), bfName.lower(), bfWidth - 1))
                            svUvmMemMapFile.write("        %s.build();\n"%(regName.lower()))
                            # (parent, size, lsb_pos, access, volatile, reset, has_reset, is_rand, individually_accessible)
                            svUvmMemMapFile.write("        %s.%s.configure(%s, %s, %s, \"RW\", 1, %s, 1, 1, 1);\n"%(regName.lower(), bfName.lower(), regName.lower(), bfWidth, regOffset, hex(bfDefault).replace("0x", "'h"))) 
                        else:
                            svUvmMemMapFile.write("        %s.configure(this, null, \"\");\n"%(regName.lower()))
                            svUvmMemMapFile.write("        %s.build();\n"%(regName.lower()))
                            for m in range(bfQueryModel.rowCount()):
                                bfRecord  = bfQueryModel.record(m)
                                bfName    = bfRecord.value("Name")
                                bfWidth   = QRegisterConst.strToInt(bfRecord.value("Width"))
                                regOffset = QRegisterConst.strToInt(bfRecord.value("RegisterOffset"))
                                bfDefault = QRegisterConst.strToInt(bfRecord.value("DefaultValue"))
                                # (parent, size, lsb_pos, access, volatile, reset, has_reset, is_rand, individually_accessible)
                                svUvmMemMapFile.write("        %s.%s.configure(%s, %s, %s, \"RW\", 1, %s, 1, 1, 1);\n"%(regName.lower(), bfName.lower(), regName.lower(), bfWidth, regOffset, hex(bfDefault).replace("0x", "'h"))) 
                                svUvmMemMapFile.write("        %s.add_hdl_path_slice(\"%s[%s:0]\", %s, %s);\n"%(regName.lower(), bfName.lower(), bfWidth - 1, regOffset, bfWidth))
                        svUvmMemMapFile.write("        default_map.add_reg(%s, %s);\n"%(regName.lower(), hex(regAddrSV).replace("0x", "'h")))
                        svUvmMemMapFile.write("\n")

                        # add reg class in regmap file, and define/create bitfield variables
                        svUvmRegMapFile.write("class %s extends uvm_reg;\n"%(regNameSV.lower()))
                        for m in range(bfQueryModel.rowCount()):
                            bfRecord  = bfQueryModel.record(m)
                            bfName    = bfRecord.value("Name")
                            svUvmRegMapFile.write("    rand uvm_reg_field %s;\n"%(bfName.lower()))
                        svUvmRegMapFile.write("    virtual function void build();\n")
                        for m in range(bfQueryModel.rowCount()):
                            bfRecord  = bfQueryModel.record(m)
                            bfName    = bfRecord.value("Name")
                            svUvmRegMapFile.write("        %s = uvm_reg_field::type_id::create(\"%s\");\n"%(bfName.lower(), bfName.lower()))
                        svUvmRegMapFile.write("    endfunction\n")
                        svUvmRegMapFile.write("    `uvm_object_utils(%s);\n"%(regNameSV.lower()))
                        svUvmRegMapFile.write("    function new(input string name = \"%s\");\n"%(regNameSV.lower()))
                        svUvmRegMapFile.write("        super.new(name, UVM_NO_COVERAGE);\n")
                        svUvmRegMapFile.write("    endfunction\n")
                        svUvmRegMapFile.write("endclass\n")
                        svUvmRegMapFile.write("\n")
                    else:
                        regArray = regMatch.string.split(':')
                        regArray0 = int(regArray[0])
                        regArray1 = int(regArray[1])
                        start = min(regArray0, regArray1)
                        end   = max(regArray0, regArray1)
                        for regI in range(start, end + 1):
                            regAddrSVI = regAddrSV + int(regWidth * (regI - start) / 8)
                            regNameI   = regName   + str(regI)
                            regNameSVI = regNameSV + str(regI)

                            # sv header
                            svHeaderLines.append("`define %s %s\n"%(regNameSVI.upper() + "_ADDR", hex(regAddrSVI).replace("0x", "'h")))
                            svHeaderLines.append("`define %s %s\n"%(regNameSVI.upper() + "_RST",  hex(regDefault).replace("0x", "'h")))

                            # uvm memmap block
                            svUvmMemMapFile.write("        %s = %s::type_id::create(\"%s\");\n"%(regNameI.lower(), regNameSVI.lower(), regNameI.lower()))
                            svUvmMemMapFile.write("        %s.configure(this, null, \"%s[%s:0]\");\n"%(regNameI.lower(), regNameI.lower(), regWidth - 1))
                            svUvmMemMapFile.write("        %s.build();\n"%(regNameI.lower()))
                            svUvmMemMapFile.write("        %s.configure(%s);\n"%(regNameI.lower(), regNameI.lower()))
                            svUvmMemMapFile.write("\n")

                            # regmap block
                            svUvmRegMapFile.write("class %s extends uvm_reg;\n"%(regNameSVI.lower()))
                            for m in range(bfQueryModel.rowCount()):
                                bfRecord  = bfQueryModel.record(m)
                                bfName    = bfRecord.value("Name")
                                svUvmRegMapFile.write("    rand uvm_reg_field %s;\n"%(bfName.lower()))
                            svUvmRegMapFile.write("    virtual function void build();\n")
                            for m in range(bfQueryModel.rowCount()):
                                bfRecord  = bfQueryModel.record(m)
                                bfName    = bfRecord.value("Name")
                                svUvmRegMapFile.write("        %s = uvm_reg_field::type_id::create(\"%s\");\n"%(bfName.lower(), bfName.lower()))
                            svUvmRegMapFile.write("    endfunction\n")
                            svUvmRegMapFile.write("    `uvm_object_utils(%s);\n"%(regNameSVI.lower()))
                            svUvmRegMapFile.write("    function new(input string name = \"%s\");\n"%(regNameSVI.lower()))
                            svUvmRegMapFile.write("        super.new(name, UVM_NO_COVERAGE);\n")
                            svUvmRegMapFile.write("    endfunction\n")
                            svUvmRegMapFile.write("endclass\n")
                            svUvmRegMapFile.write("\n")

                    # bitfield
                    for m in range(bfQueryModel.rowCount()):
                        bfRecord  = bfQueryModel.record(m)
                        bfName    = bfRecord.value("Name")
                        bfDefault = QRegisterConst.strToInt(bfRecord.value("DefaultValue"))
                        bfWidth   = QRegisterConst.strToInt(bfRecord.value("Width"))
                        regOffset = QRegisterConst.strToInt(bfRecord.value("RegisterOffset"))
                        bfMask    = ((1 << bfWidth) - 1) << regOffset
                        bfNameSV = "bf_%s_%s_%s_%s"%(memMapName, regMapName, regName, bfName)
                        svHeaderLines.append("`define %s %s\n"%(bfNameSV.upper() + "_RST", hex(bfDefault).replace("0x", "'h")))
                        svHeaderLines.append("`define %s %s\n"%(bfNameSV.upper() + "_OFS", hex(regOffset).replace("0x", "'h")))
                        svHeaderLines.append("`define %s %s\n"%(bfNameSV.upper() + "_W",   hex(bfWidth).replace("0x", "'h")))
                        svHeaderLines.append("`define %s %s\n"%(bfNameSV.upper() + "_MSK", hex(bfMask).replace("0x", "'h")))

                    # insert empty line between each register
                    svHeaderLines.append("\n")

                # memmap file
                svUvmMemMapFile.write("    endfunction\n")
                svUvmMemMapFile.write("    `uvm_object_utils(%s);\n"%(regMapNameSV.lower()))
                svUvmMemMapFile.write("    function new(input string name = \"%s\");\n"%(regMapNameSV.lower()))
                svUvmMemMapFile.write("        super.new(name, UVM_NO_COVERAGE);\n")
                svUvmMemMapFile.write("    endfunction\n")
                svUvmMemMapFile.write("endclass\n")
                svUvmMemMapFile.write("\n")

                # regmap file
                svUvmRegMapFile.write("`endif\n")
                svUvmRegMapFile.close()

            # write sv header file
            maxNameLen = 0
            for line in svHeaderLines:
                seg = line.split(' ')
                if len(seg) > 1 and len(seg[1]) > maxNameLen:
                    maxNameLen = len(seg[1])
            f = "{0} {1:<%d} {2}"%maxNameLen
            for line in svHeaderLines:
                seg = line.split(' ')
                if len(seg) == 3:
                    svHeaderFile.write(f.format(*seg))
                else:
                    svHeaderFile.write(line)

            # memmap file
            svUvmMemMapFile.write("class %s extends uvm_reg_block;\n"%(memMapNameSV.lower()))
            for j in range(regMapQueryModel.rowCount()):
                regMapRecord = regMapQueryModel.record(j)
                regMapName   = regMapRecord.value("Name")
                regMapNameSV = "rm_%s_%s"%(memMapName, regMapName)
                svUvmMemMapFile.write("    rand %s %s;\n"%(regMapNameSV.lower(), regMapName.lower()))
            svUvmMemMapFile.write("    virtual function void build();\n")
            svUvmMemMapFile.write("        default_map = create_map(\"default_map\", 0, 4, UVM_LITTLE_ENDIAN, 1);\n") # TODO: param is not meaningful for top memory map?
            for j in range(regMapQueryModel.rowCount()):
                regMapRecord = regMapQueryModel.record(j)
                regMapName   = regMapRecord.value("Name")
                regMapNameSV = "rm_%s_%s"%(memMapName, regMapName)
                svUvmMemMapFile.write("        %s = %s::type_id::create(\"%s\");\n"%(regMapName.lower(), regMapNameSV.lower(), regMapName.lower()))
                svUvmMemMapFile.write("        %s.configure(this, \"\");\n"%(regMapName.lower()))
                svUvmMemMapFile.write("        %s.build();\n"%(regMapName.lower()))
                svUvmMemMapFile.write("        default_map.add_submap(%s.default_map, 0);\n"%(regMapName.lower()))
                svUvmMemMapFile.write("        %s.lock_model();\n"%(regMapName.lower()))
                svUvmMemMapFile.write("\n")
            svUvmMemMapFile.write("    endfunction\n")
            svUvmMemMapFile.write("    `uvm_object_utils(%s);\n"%(memMapNameSV.lower()))
            svUvmMemMapFile.write("    function new(input string name = \"%s\");\n"%(memMapNameSV.lower()))
            svUvmMemMapFile.write("        super.new(name, UVM_NO_COVERAGE);\n")
            svUvmMemMapFile.write("    endfunction\n")
            svUvmMemMapFile.write("endclass\n")
            svUvmMemMapFile.write("\n")

            svHeaderFile.close()
            svUvmMemMapFile.close()
        svUvmTopFile.close()
        QMessageBox.information(parent, "Exporting verilog", "Done!", QMessageBox.Yes)
        return